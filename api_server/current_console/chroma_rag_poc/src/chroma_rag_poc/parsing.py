"""
解析模块

- 保留 Label Studio 结构化结果解析
- 增强复杂通用 JSON 的层级文本提取
- 新增 PDF / DOCX / TXT / Markdown / CSV / TSV 文档解析
"""
from __future__ import annotations

import csv
import re
from collections import OrderedDict
from io import BytesIO, StringIO
from pathlib import Path
from typing import Any, Iterable

import orjson

from .schemas import SourceRecord, TextBlock
from .text_utils import normalize_text, stable_hash

SUPPORTED_SOURCE_EXTENSIONS: dict[str, str] = {
    ".json": "JSON",
    ".jsonl": "JSON",
    ".ndjson": "JSON",
    ".ipynb": "JSON",
    ".pdf": "PDF",
    ".docx": "DOCX",
    ".pptx": "PPTX",
    ".xlsx": "Spreadsheet",
    ".xls": "Spreadsheet",
    ".png": "Image",
    ".jpg": "Image",
    ".jpeg": "Image",
    ".tif": "Image",
    ".tiff": "Image",
    ".bmp": "Image",
    ".txt": "Text",
    ".md": "Markdown",
    ".markdown": "Markdown",
    ".csv": "CSV",
    ".tsv": "TSV",
    ".log": "Log",
    ".py": "Code",
    ".js": "Code",
    ".mjs": "Code",
    ".cjs": "Code",
    ".ts": "Code",
    ".tsx": "Code",
    ".jsx": "Code",
    ".java": "Code",
    ".c": "Code",
    ".cc": "Code",
    ".cpp": "Code",
    ".cxx": "Code",
    ".h": "Code",
    ".hh": "Code",
    ".hpp": "Code",
    ".hxx": "Code",
    ".cs": "Code",
    ".go": "Code",
    ".rs": "Code",
    ".php": "Code",
    ".rb": "Code",
    ".swift": "Code",
    ".kt": "Code",
    ".kts": "Code",
    ".scala": "Code",
    ".sql": "Code",
    ".sh": "Code",
    ".bash": "Code",
    ".zsh": "Code",
    ".ps1": "Code",
    ".bat": "Code",
    ".cmd": "Code",
    ".html": "Code",
    ".htm": "Code",
    ".css": "Code",
    ".scss": "Code",
    ".sass": "Code",
    ".less": "Code",
    ".xml": "Code",
    ".yaml": "Code",
    ".yml": "Code",
    ".toml": "Code",
    ".ini": "Code",
    ".cfg": "Code",
    ".conf": "Code",
    ".properties": "Code",
    ".vue": "Code",
    ".svelte": "Code",
}

STRUCTURED_JSON_EXTENSIONS = {".json", ".ipynb"}
TABULAR_SOURCE_EXTENSIONS = {".csv", ".tsv"}
EXTERNAL_PARSER_EXTENSIONS = {
    ".pptx",
    ".xlsx",
    ".xls",
    ".png",
    ".jpg",
    ".jpeg",
    ".tif",
    ".tiff",
    ".bmp",
}
TEXT_PAYLOAD_EXTENSIONS = set(SUPPORTED_SOURCE_EXTENSIONS) - STRUCTURED_JSON_EXTENSIONS - TABULAR_SOURCE_EXTENSIONS - {".pdf", ".docx"}

PASSAGE_BLOCK_RE = re.compile(
    r"(?ms)^PASSAGE_ID:\s*(?P<passage_id>[^\n]+)\n"
    r"TITLE:\s*(?P<title>[^\n]*)\n"
    r"TEXT:\s*(?P<text>.*?)(?=^PASSAGE_ID:\s*|\Z)"
)
SENTENCE_BLOCK_RE = re.compile(
    r"(?ms)^SENTENCE_ID:\s*(?P<sentence_id>[^\n]+)\n"
    r"(?P<text>.*?)(?=^SENTENCE_ID:\s*|\Z)"
)

TITLE_KEYS = (
    "title",
    "heading",
    "name",
    "chapter",
    "section",
    "subtitle",
    "subject",
)
PREFERRED_TEXT_KEYS = (
    "title",
    "heading",
    "name",
    "text",
    "content",
    "body",
    "question",
    "answer",
    "description",
    "summary",
    "caption",
)
DOCUMENT_COLLECTION_KEYS = {
    "chunks",
    "contacts",
    "data",
    "documents",
    "docs",
    "items",
    "records",
    "entries",
    "results",
    "pages",
    "sections",
    "chapters",
    "articles",
    "rows",
    "payload",
}
LABEL_SKIP_KEYS = {
    "data",
    "items",
    "records",
    "entries",
    "results",
    "children",
    "value",
    "payload",
    "root",
}


# ============================================================
# 公共 API
# ============================================================


def supported_source_extensions() -> tuple[str, ...]:
    return tuple(SUPPORTED_SOURCE_EXTENSIONS.keys())


def is_supported_source(source_name: str) -> bool:
    return Path(source_name).suffix.lower() in SUPPORTED_SOURCE_EXTENSIONS


def get_source_kind(source_name: str) -> str:
    return SUPPORTED_SOURCE_EXTENSIONS.get(Path(source_name).suffix.lower(), "Other")



def load_source_payload(raw_bytes: bytes, source_name: str) -> list[SourceRecord]:
    """????????????"""
    suffix = Path(source_name).suffix.lower()
    if suffix in STRUCTURED_JSON_EXTENSIONS:
        return load_json_payload(raw_bytes, source_name=source_name)
    if suffix == ".pdf":
        return _load_pdf_payload(raw_bytes, source_name=source_name)
    if suffix == ".docx":
        return _load_docx_payload(raw_bytes, source_name=source_name)
    if suffix in EXTERNAL_PARSER_EXTENSIONS:
        raise ValueError(f"{suffix} requires an external parser; route it through document_intake with Docling")
    if suffix in TEXT_PAYLOAD_EXTENSIONS:
        return _load_text_payload(raw_bytes, source_name=source_name)
    if suffix in TABULAR_SOURCE_EXTENSIONS:
        return _load_tabular_payload(raw_bytes, source_name=source_name, delimiter="	" if suffix == ".tsv" else ",")
    raise ValueError(f"?????????: {suffix or 'unknown'}")


def load_source_file(source_path: str | Path) -> list[SourceRecord]:
    source_path = Path(source_path)
    return load_source_payload(source_path.read_bytes(), source_name=source_path.name)


def load_json_payload(raw_bytes: bytes, source_name: str) -> list[SourceRecord]:
    """加载单个 JSON 文件，自动检测格式并去重。"""
    payload = orjson.loads(raw_bytes)
    records = parse_payload(payload, source_name=source_name)
    return dedupe_records(records)


def load_json_file(json_path: str | Path) -> list[SourceRecord]:
    json_path = Path(json_path)
    return load_json_payload(json_path.read_bytes(), source_name=json_path.name)


def load_json_directory(json_dir: str | Path) -> list[SourceRecord]:
    json_dir = Path(json_dir)
    if not json_dir.is_dir():
        raise FileNotFoundError(f"目录不存在: {json_dir}")
    all_records: list[SourceRecord] = []
    for json_path in sorted(json_dir.glob("*.json")):
        all_records.extend(load_json_file(json_path))
    return dedupe_records(all_records)


def load_source_directory(source_dir: str | Path) -> list[SourceRecord]:
    source_dir = Path(source_dir)
    if not source_dir.is_dir():
        raise FileNotFoundError(f"目录不存在: {source_dir}")

    all_records: list[SourceRecord] = []
    for path in sorted(source_dir.rglob("*")):
        if not path.is_file() or not is_supported_source(path.name):
            continue
        all_records.extend(load_source_file(path))
    return dedupe_records(all_records)


# ============================================================
# 格式检测与解析
# ============================================================


def parse_payload(payload: Any, source_name: str) -> list[SourceRecord]:
    if _looks_like_label_studio_export(payload):
        return _parse_label_studio_tasks(payload, source_name=source_name)
    return _parse_generic_json(payload, source_name=source_name)


def dedupe_records(records: Iterable[SourceRecord]) -> list[SourceRecord]:
    seen_hashes: set[str] = set()
    unique_records: list[SourceRecord] = []
    for record in records:
        text_hash = stable_hash(record.text)
        if text_hash in seen_hashes:
            continue
        seen_hashes.add(text_hash)
        unique_records.append(record)
    return unique_records


def _looks_like_label_studio_export(payload: Any) -> bool:
    return (
        isinstance(payload, list)
        and bool(payload)
        and isinstance(payload[0], dict)
        and "annotations" in payload[0]
        and "data" in payload[0]
    )


# ============================================================
# Label Studio 解析
# ============================================================


def _parse_label_studio_tasks(tasks: list[dict[str, Any]], source_name: str) -> list[SourceRecord]:
    records: list[SourceRecord] = []

    for task_index, task in enumerate(tasks):
        results = _pick_annotation_results(task)
        task_data = task.get("data") or {}
        task_id = task.get("id", task_index)

        blocks_by_id: OrderedDict[str, dict[str, Any]] = OrderedDict()
        for order, item in enumerate(results):
            item_id = item.get("id") or f"item-{order}"
            value = item.get("value") or {}
            block = blocks_by_id.setdefault(
                item_id,
                {
                    "texts": [],
                    "labels": [],
                    "order": order,
                    "x": float(value.get("x") or 0.0),
                    "y": float(value.get("y") or 0.0),
                },
            )
            block["x"] = min(block["x"], float(value.get("x") or block["x"]))
            block["y"] = min(block["y"], float(value.get("y") or block["y"]))

            if item.get("type") == "textarea":
                for text in value.get("text") or []:
                    cleaned = normalize_text(text)
                    if cleaned:
                        block["texts"].append(cleaned)

            if item.get("type") == "rectanglelabels":
                for label in value.get("rectanglelabels") or []:
                    cleaned_label = normalize_text(label)
                    if cleaned_label:
                        block["labels"].append(cleaned_label)

        structured_blocks: list[TextBlock] = []
        page_num = _coerce_int(task_data.get("page_num")) or -1
        doc_id = _coerce_int(task_data.get("doc_id")) or -1

        for block in sorted(blocks_by_id.values(), key=lambda entry: (entry["y"], entry["x"], entry["order"])):
            text = normalize_text("\n".join(block["texts"]))
            if not text:
                continue
            block_type = block["labels"][0] if block["labels"] else "Para"
            structured_blocks.append(
                TextBlock(
                    text=text,
                    block_type=block_type,
                    order=int(block["order"]),
                    page_num=page_num,
                    doc_id=doc_id,
                    x=float(block["x"]),
                    y=float(block["y"]),
                )
            )

        text = _build_structured_text(structured_blocks)
        if not text:
            continue

        filename = normalize_text(str(task_data.get("filename") or source_name)) or source_name
        total_pages = _coerce_int(task_data.get("total_pages")) or 0

        records.append(
            SourceRecord(
                source_file=source_name,
                record_id=f"{source_name}::task-{task_id}",
                filename=filename,
                page_num=_coerce_int(task_data.get("page_num")),
                text=text,
                blocks=structured_blocks,
                metadata={
                    "doc_id": doc_id,
                    "total_pages": total_pages,
                    "source_kind": "JSON",
                },
                total_pages=total_pages,
                doc_id=doc_id,
            )
        )

    return records


def _pick_annotation_results(task: dict[str, Any]) -> list[dict[str, Any]]:
    annotations = task.get("annotations") or []
    if annotations:
        best = max(annotations, key=lambda item: item.get("result_count", 0))
        return best.get("result") or []
    predictions = task.get("predictions") or []
    if predictions:
        best = max(predictions, key=lambda item: item.get("result_count", 0))
        return best.get("result") or []
    return []


def _build_structured_text(blocks: list[TextBlock]) -> str:
    parts: list[str] = []
    for block in blocks:
        cleaned = normalize_text(block.text)
        if not cleaned:
            continue
        if block.block_type.lower() == "list":
            parts.append(f"- {cleaned}")
        else:
            parts.append(cleaned)
    return "\n\n".join(parts)


# ============================================================
# 通用 JSON 解析
# ============================================================


def _parse_generic_json(payload: Any, source_name: str) -> list[SourceRecord]:
    candidate_nodes = _discover_document_nodes(payload)
    records = [
        record
        for index, (node, context) in enumerate(candidate_nodes)
        if (record := _record_from_generic_node(node, source_name, index, context)) is not None
    ]
    if records:
        return records

    blocks = _collect_structured_blocks(payload)
    text = _build_structured_text(blocks)
    if not text:
        return []
    return [
        SourceRecord(
            source_file=source_name,
            record_id=f"{source_name}::root",
            filename=source_name,
            page_num=None,
            text=text,
            blocks=blocks,
            metadata={"source_kind": "JSON"},
        )
    ]


def _discover_document_nodes(payload: Any) -> list[tuple[Any, tuple[str, ...]]]:
    if isinstance(payload, list):
        if len(payload) <= 1:
            return [(item, ("entry",)) for item in payload]
        return [(item, (f"entry {index + 1}",)) for index, item in enumerate(payload)]

    if isinstance(payload, dict):
        nested = _find_document_collection(payload)
        if nested:
            prefix, items = nested
            return [(item, prefix + (f"entry {index + 1}",)) for index, item in enumerate(items)]
        return [(payload, tuple())]

    return [(payload, tuple())]


def _find_document_collection(node: dict[str, Any], path: tuple[str, ...] = ()) -> tuple[tuple[str, ...], list[Any]] | None:
    for key, value in node.items():
        normalized_key = str(key).strip().lower()
        if normalized_key in DOCUMENT_COLLECTION_KEYS and isinstance(value, list) and value:
            if any(isinstance(item, (dict, list, str)) for item in value):
                return path + (str(key),), value
        if isinstance(value, dict):
            found = _find_document_collection(value, path + (str(key),))
            if found:
                return found
    return None


def _record_from_generic_node(
    node: Any,
    source_name: str,
    index: int,
    context: tuple[str, ...],
) -> SourceRecord | None:
    blocks = _collect_structured_blocks(node, context=context)
    text = _build_structured_text(blocks)
    if not text:
        return None

    identifier = _extract_identifier(node, index)
    page_num = _extract_page_num(node)
    filename = _extract_title(node) or _path_label(context) or source_name
    return SourceRecord(
        source_file=source_name,
        record_id=f"{source_name}::item-{identifier}",
        filename=filename,
        page_num=page_num,
        text=text,
        blocks=blocks,
        metadata={"source_kind": "JSON", **_extract_scalar_metadata(node)},
    )


def _collect_structured_blocks(
    node: Any,
    *,
    context: tuple[str, ...] = (),
    page_num: int | None = None,
) -> list[TextBlock]:
    blocks: list[TextBlock] = []

    def append_block(text: str, block_type: str = "Para") -> None:
        cleaned = normalize_text(text)
        if not cleaned:
            return
        if blocks and blocks[-1].text == cleaned and blocks[-1].block_type == block_type:
            return
        blocks.append(
            TextBlock(
                text=cleaned,
                block_type=block_type,
                order=len(blocks),
                page_num=page_num if page_num is not None else -1,
            )
        )

    def visit(value: Any, path: tuple[str, ...]) -> None:
        if isinstance(value, dict):
            heading = _extract_title(value)
            if heading:
                append_block(heading, "Title")

            ordered_keys = _ordered_keys(value)
            for key in ordered_keys:
                nested_value = value[key]
                if heading and str(key).strip().lower() in TITLE_KEYS and normalize_text(str(nested_value)) == heading:
                    continue
                visit(nested_value, path + (str(key),))
            return

        if isinstance(value, list):
            for index, item in enumerate(value):
                if isinstance(item, (dict, list)):
                    visit(item, path + (f"item {index + 1}",))
                else:
                    formatted = _format_scalar_fragment(item, path)
                    if formatted:
                        append_block(formatted, "List" if len(value) > 1 else "Para")
            return

        formatted = _format_scalar_fragment(value, path)
        if formatted:
            append_block(formatted, "Para")

    visit(node, context)
    return blocks


def _ordered_keys(node: dict[str, Any]) -> list[str]:
    ordered: list[str] = []
    for preferred in PREFERRED_TEXT_KEYS:
        for key in node.keys():
            if str(key).strip().lower() == preferred and key not in ordered:
                ordered.append(str(key))
    for key in node.keys():
        key_str = str(key)
        if key_str not in ordered:
            ordered.append(key_str)
    return ordered


def _format_scalar_fragment(value: Any, path: tuple[str, ...]) -> str:
    if value is None:
        return ""
    if isinstance(value, str):
        cleaned = normalize_text(value)
        if len(cleaned) < 2:
            return ""
    elif isinstance(value, (int, float, bool)):
        cleaned = normalize_text(str(value))
    else:
        return ""

    label = _path_label(path)
    if label:
        return f"{label}: {cleaned}"
    return cleaned


def _path_label(path: tuple[str, ...]) -> str:
    labels: list[str] = []
    for item in path:
        cleaned = normalize_text(str(item))
        if not cleaned:
            continue
        lowered = cleaned.lower()
        if lowered in LABEL_SKIP_KEYS:
            continue
        if lowered.startswith("entry ") or lowered.startswith("item "):
            continue
        labels.append(cleaned)
    if not labels:
        return ""
    return " / ".join(labels[-2:])


def _extract_title(node: Any) -> str:
    if not isinstance(node, dict):
        return ""
    for key in TITLE_KEYS:
        for actual_key, value in node.items():
            if str(actual_key).strip().lower() != key:
                continue
            if isinstance(value, str):
                cleaned = normalize_text(value)
                if cleaned:
                    return cleaned
    return ""


def _extract_identifier(node: Any, fallback: int) -> str:
    if isinstance(node, dict):
        for key in ("id", "uuid", "record_id", "code", "name", "title"):
            value = node.get(key)
            if value is None:
                continue
            cleaned = normalize_text(str(value))
            if cleaned:
                return cleaned
    return str(fallback)


def _extract_page_num(node: Any) -> int | None:
    if not isinstance(node, dict):
        return None
    for key in ("page_num", "page", "pageIndex", "page_index"):
        if key in node:
            return _coerce_int(node.get(key))
    return None


def _extract_scalar_metadata(node: Any) -> dict[str, str | int | float | bool]:
    if not isinstance(node, dict):
        return {}
    excluded = {"text", "content", "body", "attachments", "messages", "chunks"}
    metadata: dict[str, str | int | float | bool] = {}
    for key, value in node.items():
        key_str = str(key).strip()
        if not key_str or key_str in excluded:
            continue
        if isinstance(value, (str, int, float, bool)):
            metadata[key_str] = value
    return metadata


# ============================================================
# 文档文件解析
# ============================================================


def _load_pdf_payload(raw_bytes: bytes, source_name: str) -> list[SourceRecord]:
    try:
        from pypdf import PdfReader
    except ImportError as exc:
        raise ValueError("PDF 解析依赖缺失，请安装 pypdf") from exc

    reader = PdfReader(BytesIO(raw_bytes))
    total_pages = len(reader.pages)
    records: list[SourceRecord] = []

    for page_index, page in enumerate(reader.pages, start=1):
        text = normalize_text(page.extract_text() or "")
        if not text:
            continue
        blocks = _paragraph_blocks(text, page_num=page_index)
        records.append(
            SourceRecord(
                source_file=source_name,
                record_id=f"{source_name}::page-{page_index}",
                filename=source_name,
                page_num=page_index,
                text=_build_structured_text(blocks),
                blocks=blocks,
                metadata={"source_kind": "PDF"},
                total_pages=total_pages,
            )
        )

    if not records:
        raise ValueError("PDF 未提取到可用文本，可能是扫描件、加密文档或图片型 PDF")
    return records


def _load_docx_payload(raw_bytes: bytes, source_name: str) -> list[SourceRecord]:
    try:
        from docx import Document
    except ImportError as exc:
        raise ValueError("DOCX 解析依赖缺失，请安装 python-docx") from exc

    document = Document(BytesIO(raw_bytes))
    blocks: list[TextBlock] = []

    def append(text: str, block_type: str = "Para") -> None:
        cleaned = normalize_text(text)
        if not cleaned:
            return
        blocks.append(TextBlock(text=cleaned, block_type=block_type, order=len(blocks)))

    for paragraph in document.paragraphs:
        text = normalize_text(paragraph.text)
        if not text:
            continue
        style_name = normalize_text(getattr(getattr(paragraph, "style", None), "name", "")).lower()
        block_type = "Title" if any(token in style_name for token in ("heading", "title")) else "Para"
        append(text, block_type)

    for table in document.tables:
        for row in table.rows:
            cells = [normalize_text(cell.text) for cell in row.cells if normalize_text(cell.text)]
            if cells:
                append(" | ".join(cells), "List")

    if not blocks:
        raise ValueError("DOCX 未提取到可用文本")

    return [
        SourceRecord(
            source_file=source_name,
            record_id=f"{source_name}::doc",
            filename=source_name,
            page_num=None,
            text=_build_structured_text(blocks),
            blocks=blocks,
            metadata={"source_kind": "DOCX"},
        )
    ]


def _load_text_payload(raw_bytes: bytes, source_name: str) -> list[SourceRecord]:
    text = _decode_text_bytes(raw_bytes)
    passage_records = _load_passage_delimited_text(text, source_name)
    if passage_records:
        return passage_records
    sentence_records = _load_sentence_delimited_text(text, source_name)
    if sentence_records:
        return sentence_records
    blocks = _paragraph_blocks(text)
    if not blocks:
        raise ValueError("文本文件为空或未提取到可用内容")
    return [
        SourceRecord(
            source_file=source_name,
            record_id=f"{source_name}::text",
            filename=source_name,
            page_num=None,
            text=_build_structured_text(blocks),
            blocks=blocks,
            metadata={"source_kind": get_source_kind(source_name)},
        )
    ]


def _load_passage_delimited_text(text: str, source_name: str) -> list[SourceRecord]:
    records: list[SourceRecord] = []
    for index, match in enumerate(PASSAGE_BLOCK_RE.finditer(text), start=1):
        passage_id = normalize_text(match.group("passage_id"))
        body = normalize_text(match.group("text"))
        if not passage_id or not body:
            continue
        title = normalize_text(match.group("title"))
        section_id, section_path = _section_from_passage_id(passage_id)
        blocks = [
            TextBlock(text=f"PASSAGE_ID: {passage_id}", block_type="Title", order=0),
        ]
        if title:
            blocks.append(TextBlock(text=f"TITLE: {title}", block_type="Title", order=len(blocks)))
        blocks.append(TextBlock(text=body, block_type="Para", order=len(blocks)))
        records.append(
            SourceRecord(
                source_file=source_name,
                record_id=f"{source_name}::passage-{passage_id}",
                filename=source_name,
                page_num=None,
                text=_build_structured_text(blocks),
                blocks=blocks,
                metadata={
                    "source_kind": get_source_kind(source_name),
                    "boundary_type": "passage",
                    "passage_id": passage_id,
                    "passage_title": title,
                    "section_id": section_id,
                    "section_path": section_path,
                    "passage_ordinal": index,
                },
            )
        )
    return records


def _load_sentence_delimited_text(text: str, source_name: str) -> list[SourceRecord]:
    if "SENTENCE_ID:" not in text:
        return []
    ragbench_id = _first_header_value(text, "RAGBENCH_ID")
    records: list[SourceRecord] = []
    for index, match in enumerate(SENTENCE_BLOCK_RE.finditer(text), start=1):
        sentence_id = normalize_text(match.group("sentence_id"))
        sentence_text = normalize_text(match.group("text"))
        if not sentence_id or not sentence_text:
            continue
        blocks = [
            TextBlock(text=f"SENTENCE_ID: {sentence_id}", block_type="Title", order=0),
            TextBlock(text=sentence_text, block_type="Para", order=1),
        ]
        document_id = sentence_id.split("a", 1)[0] if sentence_id[:1].isdigit() else sentence_id.split(".", 1)[0]
        records.append(
            SourceRecord(
                source_file=source_name,
                record_id=f"{source_name}::sentence-{sentence_id}",
                filename=source_name,
                page_num=None,
                text=_build_structured_text(blocks),
                blocks=blocks,
                metadata={
                    "source_kind": get_source_kind(source_name),
                    "boundary_type": "sentence",
                    "sentence_id": sentence_id,
                    "document_id": document_id,
                    "ragbench_id": ragbench_id,
                    "sentence_ordinal": index,
                },
            )
        )
    return records


def _first_header_value(text: str, header: str) -> str:
    match = re.search(rf"(?m)^{re.escape(header)}:\s*(.+)$", text)
    return normalize_text(match.group(1)) if match else ""


def _section_from_passage_id(passage_id: str) -> tuple[str, str]:
    parts = [part for part in re.split(r"[-/]+", passage_id) if part]
    if len(parts) <= 1:
        return passage_id, passage_id
    section_id = "-".join(parts[:-1])
    section_path = " > ".join(parts)
    return section_id, section_path


def _load_tabular_payload(raw_bytes: bytes, source_name: str, delimiter: str) -> list[SourceRecord]:
    text = _decode_text_bytes(raw_bytes)
    reader = csv.reader(StringIO(text), delimiter=delimiter)
    rows = [[normalize_text(cell) for cell in row] for row in reader]
    rows = [[cell for cell in row if cell] for row in rows if any(cell for cell in row)]
    if not rows:
        raise ValueError("表格文件为空")

    header = rows[0]
    blocks: list[TextBlock] = []

    def append(text: str, block_type: str = "List") -> None:
        cleaned = normalize_text(text)
        if cleaned:
            blocks.append(TextBlock(text=cleaned, block_type=block_type, order=len(blocks)))

    append(f"Columns: {' | '.join(header)}", "Title")
    for index, row in enumerate(rows[1:], start=1):
        pairs: list[str] = []
        for col_index, cell in enumerate(row):
            label = header[col_index] if col_index < len(header) else f"Column {col_index + 1}"
            pairs.append(f"{label}: {cell}")
        append(f"Row {index}: " + " | ".join(pairs))

    return [
        SourceRecord(
            source_file=source_name,
            record_id=f"{source_name}::table",
            filename=source_name,
            page_num=None,
            text=_build_structured_text(blocks),
            blocks=blocks,
            metadata={"source_kind": get_source_kind(source_name)},
        )
    ]


def _decode_text_bytes(raw_bytes: bytes) -> str:
    for encoding in ("utf-8-sig", "utf-8", "utf-16", "gb18030"):
        try:
            return raw_bytes.decode(encoding)
        except UnicodeDecodeError:
            continue
    return raw_bytes.decode("utf-8", errors="ignore")


def _paragraph_blocks(text: str, *, page_num: int | None = None) -> list[TextBlock]:
    paragraphs = [normalize_text(part) for part in text.replace("\r\n", "\n").split("\n\n")]
    paragraphs = [part for part in paragraphs if part]
    if not paragraphs:
        merged = normalize_text(text)
        if not merged:
            return []
        paragraphs = [merged]

    blocks: list[TextBlock] = []
    for index, paragraph in enumerate(paragraphs):
        block_type = "Title" if index == 0 and len(paragraph) <= 80 else "Para"
        blocks.append(
            TextBlock(
                text=paragraph,
                block_type=block_type,
                order=index,
                page_num=page_num if page_num is not None else -1,
            )
        )
    return blocks


# ============================================================
# 辅助函数
# ============================================================


def _dedupe_preserve_order(values: list[str]) -> list[str]:
    seen: set[str] = set()
    output: list[str] = []
    for value in values:
        if value in seen:
            continue
        seen.add(value)
        output.append(value)
    return output


def _coerce_int(value: Any) -> int | None:
    try:
        return int(value) if value is not None else None
    except (TypeError, ValueError):
        return None
