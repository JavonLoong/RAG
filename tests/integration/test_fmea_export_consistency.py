from __future__ import annotations

# ruff: noqa: RUF001
import io
import json
from collections.abc import Iterable
from dataclasses import fields, replace
from html import unescape
from zipfile import ZipFile

import openpyxl
import orjson
import pytest
from defusedxml.ElementTree import fromstring as safe_xml_fromstring
from docx import Document
from docx.oxml.ns import qn
from docx.table import Table
from docx.text.paragraph import Paragraph

from fmea_application.snapshot_contracts import NormalizedFmeaSnapshot, build_normalized_snapshot, snapshot_content_hash
from fmea_infrastructure.export_docx import DocxFmeaExporter
from fmea_infrastructure.export_json import CanonicalJsonExporter
from fmea_infrastructure.export_xlsx import XlsxFmeaExporter
from tests.fmea_governance_fixtures import make_fmea_revision, make_normalized_snapshot, make_readiness_issue
from tests.unit.test_fmea_snapshot_contracts import _marked_publication_body_source

MARKER = "DRAFT PREVIEW — NOT PUBLISHED"
FORMAT_MEDIA_TYPES = {
    "json": "application/json",
    "xlsx": "application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
    "docx": "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
}
ENVELOPE_KEYS = (
    "schema_version",
    "snapshot_schema_version",
    "snapshot_id",
    "workspace_id",
    "analysis_id",
    "revision_id",
    "revision_hash",
    "publication_id",
    "source_publication_id",
    "manifest_id",
    "row_count",
    "snapshot_hash",
    "created_at",
    "draft_preview",
    "draft_marker",
    "format",
    "media_type",
)


def _decode_cell(value: object, value_type: str) -> object:
    if value_type == "str":
        return "" if value is None else str(value)
    if value_type == "null":
        return None
    if value_type == "bool":
        return json.loads(str(value).lower())
    if value_type in {"int", "float", "json"}:
        return json.loads(str(value))
    raise AssertionError


def _parse_typed_table(rows: Iterable[tuple[object, ...]]) -> list[dict[str, object]]:
    rows = iter(rows)
    headers = [str(value) for value in next(rows)]
    type_column = headers.index("__types__")
    result: list[dict[str, object]] = []
    for row in rows:
        if not any(value is not None and value != "" for value in row):
            continue
        types = json.loads(str(row[type_column]))
        item: dict[str, object] = {}
        for header, value in zip(headers, row, strict=True):
            if header in {"Identity", "__types__"} or header not in types:
                continue
            item[header] = _decode_cell(value, types[header])
        result.append(item)
    return result


def _parse_xlsx(payload: bytes) -> dict[str, object]:
    workbook = openpyxl.load_workbook(io.BytesIO(payload), data_only=False, read_only=False)
    manifest_rows = list(workbook["Manifest"].iter_rows(values_only=True))
    manifest = {str(row[0]): _decode_cell(row[1], str(row[2])) for row in manifest_rows[1:] if row[0] is not None}
    result: dict[str, object] = {key: manifest[key] for key in ENVELOPE_KEYS}
    result.update({
        "rows": _parse_typed_table(workbook["FMEA"].iter_rows(values_only=True)),
        "risk_records": _parse_typed_table(workbook["Risk"].iter_rows(values_only=True)),
        "evidence_summary": _parse_typed_table(workbook["Evidence"].iter_rows(values_only=True)),
        "decision_summary": _parse_typed_table(workbook["Decisions"].iter_rows(values_only=True)),
        "unresolved_items": _parse_typed_table(workbook["Unresolved"].iter_rows(values_only=True)),
        "version_manifest": manifest["version_manifest"],
        "audit_summary": manifest["audit_summary"],
    })
    propagation_rows = _parse_typed_table(workbook["Propagation"].iter_rows(values_only=True))
    result["propagation"] = propagation_rows[0] if propagation_rows else None
    return result


def _parse_docx_typed_table(table) -> list[dict[str, object]]:
    return _parse_typed_table(tuple(tuple(cell.text for cell in row.cells) for row in table.rows))


def _canonical_docx_tables(document: Document) -> dict[str, Table]:
    tables: dict[str, Table] = {}
    pending_name: str | None = None
    for element in document.element.body.iterchildren():
        if element.tag == qn("w:p"):
            paragraph_text = Paragraph(element, document).text
            prefix = "Canonical table: "
            pending_name = paragraph_text.removeprefix(prefix) if paragraph_text.startswith(prefix) else None
        elif element.tag == qn("w:tbl") and pending_name is not None:
            tables[pending_name] = Table(element, document)
            pending_name = None
    return tables


def _parse_docx(payload: bytes) -> dict[str, object]:
    document = Document(io.BytesIO(payload))
    tables = _canonical_docx_tables(document)
    assert set(tables) >= {"Manifest", "FMEA", "Risk", "Propagation", "Evidence", "Decisions", "Unresolved"}
    manifest_table = tables["Manifest"]
    manifest = {
        row.cells[0].text: _decode_cell(row.cells[1].text, row.cells[2].text) for row in manifest_table.rows[1:]
    }
    result: dict[str, object] = {key: manifest[key] for key in ENVELOPE_KEYS}
    result.update({
        "rows": _parse_docx_typed_table(tables["FMEA"]),
        "risk_records": _parse_docx_typed_table(tables["Risk"]),
        "propagation": None,
    })
    propagation_rows = _parse_docx_typed_table(tables["Propagation"])
    result["propagation"] = propagation_rows[0] if propagation_rows else None
    result.update({
        "evidence_summary": _parse_docx_typed_table(tables["Evidence"]),
        "decision_summary": _parse_docx_typed_table(tables["Decisions"]),
        "unresolved_items": _parse_docx_typed_table(tables["Unresolved"]),
        "version_manifest": manifest["version_manifest"],
        "audit_summary": manifest["audit_summary"],
    })
    return result


def _semantic_json(snapshot, *, draft_preview: bool = False) -> dict[str, object]:
    body = orjson.loads(CanonicalJsonExporter().render(snapshot, draft_preview=draft_preview))
    return body


def _without_format_identity(view: dict[str, object]) -> dict[str, object]:
    return {key: value for key, value in view.items() if key not in {"format", "media_type"}}


def _assert_format_identity(view: dict[str, object], export_format: str) -> None:
    assert view["format"] == export_format
    assert view["media_type"] == FORMAT_MEDIA_TYPES[export_format]


def _forge_snapshot(snapshot: NormalizedFmeaSnapshot, **overrides: object) -> NormalizedFmeaSnapshot:
    forged = object.__new__(NormalizedFmeaSnapshot)
    for field in fields(snapshot):
        object.__setattr__(forged, field.name, getattr(snapshot, field.name))
    for field_name, value in overrides.items():
        object.__setattr__(forged, field_name, value)
    return forged


def _readable_snapshot() -> NormalizedFmeaSnapshot:
    source = _marked_publication_body_source()
    row = dict(source.rows[0])
    row.update({
        "item_id": "燃料过滤器",
        "function_id": "稳定供油",
        "failure_mode": "燃料滤清器堵塞",
        "causes": ("杂质积聚", "低温结蜡", "维护周期过长"),
        "effects": ("供油压力下降", "燃烧不稳定"),
        "controls": ("压差监测",),
        "actions": ("清洁或更换滤芯",),
        "field_evidence": (
            {"field_key": "failure_mode", "evidence_ids": ("evidence-1",)},
            {"field_key": "causes", "evidence_ids": ("evidence-1",)},
        ),
        "field_support": (
            {"field_key": "failure_mode", "support_status": "supported"},
            {"field_key": "causes", "support_status": "supported"},
        ),
        "extension_values": (
            {"field_key": "fuel.pressure_drop", "value_type": "decimal", "value": "48.2000"},
            {"field_key": "vendor.unrecognized", "value_type": "string", "value": "keep this exact"},
        ),
    })
    evidence_ref = {
        "evidence_id": "evidence-1",
        "document_id": "manual-1",
        "document_version": "7",
        "content_hash": "c" * 64,
        "evidence_hash": "d" * 64,
        "locator": {"page": 12, "span": 3},
        "quote": "滤清器堵塞会造成燃料压力下降，维护记录应核对更换周期。" * 8,
        "source_type": "primary_document",
        "source_trust": "trusted",
    }
    risk = {
        "assessment_id": "assessment-1",
        "assessment_hash": "e" * 64,
        "workspace_id": source.publication_workspace_id,
        "row_id": "row-1",
        "source_record_version": 1,
        "evidence_pack_id": "pack-1",
        "domain_pack_id": "fuel-domain",
        "domain_pack_version": "1.0.0",
        "rule_pack_id": "fuel-sod-rpn",
        "rule_pack_version": "1.0.0",
        "status": "confirmed",
        "dimensions": (
            {
                "name": "severity",
                "value": 3,
                "evidence_ids": ("evidence-1",),
                "reason": "原生确认记录",
                "uncertainty": None,
            },
        ),
        "derived": {"rpn": 12, "evidence_ids": ("evidence-1",)},
        "proposal_id": None,
        "invalidated_reason": None,
        "record_version": 1,
        "confirmation_basis": None,
    }
    decision = dict(source.decision_summary[0])
    decision.update({"row_id": "row-1", "record_version": 1, "row_hash": "a" * 64})
    layout = {
        "template_identity": {"template_id": "fuel-fmea", "version": "1.0.0", "template_hash": "a" * 64},
        "columns": [
            {"field_key": "failure_mode", "label": "故障模式", "value_type": "string", "value_path": ("row", "failure_mode")},
            {"field_key": "causes", "label": "原因", "value_type": "string[]", "value_path": ("row", "causes")},
            {"field_key": "effects", "label": "影响", "value_type": "string[]", "value_path": ("row", "effects")},
            {"field_key": "actions", "label": "措施", "value_type": "string[]", "value_path": ("row", "actions")},
            {
                "field_key": "fuel.pressure_drop",
                "label": "压降",
                "value_type": "decimal",
                "value_path": ("extension_values", "fuel.pressure_drop"),
            },
        ],
    }
    return build_normalized_snapshot(
        replace(
            source,
            rows=(row,),
            risk_records=(risk,),
            evidence_summary=({**source.evidence_summary[0], "refs": (evidence_ref,)},),
            decision_summary=(decision,),
            version_manifest={
                **source.version_manifest,
                "body_schema_version": "graphrag.fmea.body.v1",
                "template_identities": (("fuel-fmea", "1.0.0", "a" * 64),),
                "report_layout": layout,
            },
        )
    )


def _scan_office_package(payload: bytes) -> tuple[tuple[str, ...], dict[str, int]]:
    marker_encodings = tuple(MARKER.encode(encoding) for encoding in ("utf-8", "utf-16-le", "utf-16-be"))
    forbidden_parts = ("altchunk", "externallinks", "vbaproject", "embeddings")
    executable_suffixes = (".bin", ".com", ".dll", ".exe", ".jar", ".js", ".ps1", ".vbs")
    marker_hits: dict[str, int] = {}
    with ZipFile(io.BytesIO(payload)) as archive:
        names = tuple(archive.namelist())
        for name in names:
            folded_name = name.casefold()
            assert not name.startswith(("/", "\\"))
            assert "\\" not in name
            assert ".." not in name.split("/")
            assert not any(part in folded_name for part in forbidden_parts)
            assert not folded_name.endswith(executable_suffixes)

            payload_part = archive.read(name)
            count = max(payload_part.count(encoded) for encoded in marker_encodings)
            if folded_name.endswith((".xml", ".rels")):
                decoded_xml = unescape(payload_part.decode("utf-8"))
                root = safe_xml_fromstring(payload_part)
                visible_text = "".join(root.itertext())
                count = max(count, decoded_xml.count(MARKER), visible_text.count(MARKER))
                for element in root.iter():
                    local_name = element.tag.rsplit("}", 1)[-1].casefold()
                    assert local_name not in {"altchunk", "fldchar", "fldsimple", "instrtext"}
                    if folded_name.endswith(".rels"):
                        target_mode = element.attrib.get("TargetMode", "")
                        target = element.attrib.get("Target", "").casefold()
                        assert target_mode.casefold() != "external"
                        assert not target.startswith(("file:", "ftp:", "http:", "https:"))
            if count:
                marker_hits[name] = count
        return names, marker_hits


def _snapshot_with_reserved_marker(location: str) -> NormalizedFmeaSnapshot:
    collision = f"confidential-before {MARKER} confidential-after"
    snapshot = make_normalized_snapshot()
    overrides: dict[str, object] = {
        "rows": ({"row_id": "row-1", "nested": {"value": collision}},),
        "risk_records": ({"assessment_id": "assessment-1", "nested": [collision]},),
        "propagation": {"nested": {"value": collision}},
        "evidence_summary": ({"pack_id": "pack-1", "nested": [collision]},),
        "decision_summary": ({"decision_id": "decision-1", "nested": {"value": collision}},),
        "unresolved_items": ({"code": "collision", "nested": [collision]},),
        "version_manifest": {"nested": {"value": collision}},
        "audit_summary": {"nested": [{"value": collision}]},
    }
    forged = _forge_snapshot(snapshot, **{location: overrides[location]})
    if location == "rows":
        object.__setattr__(forged, "row_count", 1)
    object.__setattr__(forged, "snapshot_hash", snapshot_content_hash(forged))
    return forged


def _docx_footer_identity(payload: bytes) -> dict[str, str]:
    document = Document(io.BytesIO(payload))
    footer = " | ".join(paragraph.text for paragraph in document.sections[0].footer.paragraphs)
    return dict(part.split("=", 1) for part in footer.split(" | ") if "=" in part)


def test_json_xlsx_docx_share_all_snapshot_semantics() -> None:
    revision = make_fmea_revision(
        unresolved_items=(make_readiness_issue(code="missing-review", severity="blocking", evidence_ids=("ev-1",)),)
    )
    snapshot = make_normalized_snapshot(
        revision=revision,
        rows=2,
        risk_records=(
            {"assessment_id": "assessment-1", "status": "confirmed", "rpn": 12},
            {"assessment_id": "assessment-2", "status": "proposed", "rpn": 6},
        ),
        propagation={"graph_revision_id": "graph-1", "edges": [{"edge_id": "edge-1", "evidence_ids": ["ev-1"]}]},
        evidence_summary=({"pack_id": "pack-1", "evidence_count": 2, "evidence_ids": ["ev-1", "ev-2"]},),
        decision_summary=({"decision_id": "decision-1", "action": "accept", "actor": "reviewer-1"},),
        version_manifest={
            "schema_id": "graphrag.fmea.v1",
            "domain_pack": {"id": "fuel-combustion", "version": "1.0.0", "hash": "a" * 64},
            "template": {"id": "fuel-fmea", "version": "1.0.0", "hash": "b" * 64},
        },
        audit_summary={"event_count": 3, "last_hash": "c" * 64},
    )

    json_view = _semantic_json(snapshot)
    xlsx_view = _parse_xlsx(XlsxFmeaExporter().render(snapshot))
    docx_view = _parse_docx(DocxFmeaExporter().render(snapshot))

    assert _without_format_identity(json_view) == _without_format_identity(xlsx_view)
    assert _without_format_identity(json_view) == _without_format_identity(docx_view)
    _assert_format_identity(json_view, "json")
    _assert_format_identity(xlsx_view, "xlsx")
    _assert_format_identity(docx_view, "docx")


def test_readable_office_bodies_keep_real_content_while_machine_decoders_remain_lossless() -> None:
    snapshot = _readable_snapshot()
    expected = _semantic_json(snapshot)
    xlsx_payload = XlsxFmeaExporter().render(snapshot)
    docx_payload = DocxFmeaExporter().render(snapshot)

    assert _without_format_identity(_parse_xlsx(xlsx_payload)) == _without_format_identity(expected)
    assert _without_format_identity(_parse_docx(docx_payload)) == _without_format_identity(expected)

    workbook = openpyxl.load_workbook(io.BytesIO(xlsx_payload), data_only=False, read_only=False)
    body_text = "\n".join(str(value) for row in workbook["正文"].iter_rows(values_only=True) for value in row if value is not None)
    detail_text = "\n".join(
        str(value) for row in workbook["正文详情"].iter_rows(values_only=True) for value in row if value is not None
    )
    docx = Document(io.BytesIO(docx_payload))
    docx_text = "\n".join(paragraph.text for paragraph in docx.paragraphs) + "\n" + "\n".join(
        cell.text for table in docx.tables for row in table.rows for cell in row.cells
    )

    quote = str(snapshot.evidence_summary[0]["refs"][0]["quote"])
    assert "燃料滤清器堵塞" in body_text
    assert "48.2000" in body_text
    assert "vendor.unrecognized" in detail_text
    assert "keep this exact" in detail_text
    assert quote in detail_text
    assert detail_text.count(quote) == 1
    assert "evidence-1.quote" in detail_text
    assert "evidence-1.document_id" in detail_text
    assert "assessment-1" in detail_text
    assert "assessment-1.assessment_id" in detail_text
    assert "assessment-1.derived.rpn" in detail_text
    assert "decision-1.decision" in detail_text
    assert "复核" in detail_text
    assert "燃料滤清器堵塞" in docx_text
    assert "48.2000" in docx_text
    assert "vendor.unrecognized" in docx_text
    assert quote in docx_text
    assert "evidence-1.quote" in docx_text
    assert "evidence-1.document_id" in docx_text
    docx_reading_text = "\n".join(paragraph.text for paragraph in docx.paragraphs)
    assert docx_reading_text.count(quote) == 1
    assert "assessment-1" in docx_text
    assert "assessment-1.assessment_id" in docx_text
    assert "assessment-1.derived.rpn" in docx_text
    assert "decision-1.decision" in docx_text


def test_empty_optional_parts_round_trip_without_fabrication() -> None:
    snapshot = make_normalized_snapshot(
        revision=make_fmea_revision(unresolved_items=()),
        risk_records=(),
        propagation=None,
        evidence_summary=(),
        decision_summary=(),
    )

    expected = _semantic_json(snapshot)
    assert _without_format_identity(_parse_xlsx(XlsxFmeaExporter().render(snapshot))) == _without_format_identity(
        expected
    )
    assert _without_format_identity(_parse_docx(DocxFmeaExporter().render(snapshot))) == _without_format_identity(
        expected
    )


def test_empty_mapping_record_round_trips_through_xlsx_types_marker() -> None:
    snapshot = _forge_snapshot(make_normalized_snapshot(), unresolved_items=({},))
    snapshot = _forge_snapshot(snapshot, snapshot_hash=snapshot_content_hash(snapshot))

    expected = _semantic_json(snapshot)
    payload = XlsxFmeaExporter().render(snapshot)

    assert _without_format_identity(_parse_xlsx(payload)) == _without_format_identity(expected)
    workbook = openpyxl.load_workbook(io.BytesIO(payload), data_only=False, read_only=False)
    try:
        unresolved_rows = list(workbook["Unresolved"].iter_rows(values_only=True))
        assert unresolved_rows[0] == ("Identity", "__types__")
        assert unresolved_rows[1] == ("item-001", "{}")
    finally:
        workbook.close()


def test_draft_marker_is_visible_in_all_formats_and_absent_from_published() -> None:
    snapshot = make_normalized_snapshot()
    rendered = {
        "json": CanonicalJsonExporter().render(snapshot, draft_preview=True),
        "xlsx": XlsxFmeaExporter().render(snapshot, draft_preview=True),
        "docx": DocxFmeaExporter().render(snapshot, draft_preview=True),
    }
    parsed = {
        "json": orjson.loads(rendered["json"]),
        "xlsx": _parse_xlsx(rendered["xlsx"]),
        "docx": _parse_docx(rendered["docx"]),
    }
    assert _without_format_identity(parsed["json"]) == _without_format_identity(parsed["xlsx"])
    assert _without_format_identity(parsed["json"]) == _without_format_identity(parsed["docx"])
    for export_format, view in parsed.items():
        _assert_format_identity(view, export_format)
        assert view["draft_preview"] is True
        assert view["draft_marker"] == MARKER
        assert view["publication_id"] is None
        assert view["source_publication_id"] == snapshot.publication_id

    footer = _docx_footer_identity(rendered["docx"])
    assert footer == {
        "revision_id": snapshot.revision_id,
        "snapshot_id": snapshot.snapshot_id,
        "publication_id": "",
        "source_publication_id": snapshot.publication_id,
        "snapshot_hash": snapshot.snapshot_hash,
    }
    assert MARKER in rendered["json"].decode("utf-8")

    published = {
        "json": CanonicalJsonExporter().render(snapshot, draft_preview=False),
        "xlsx": XlsxFmeaExporter().render(snapshot, draft_preview=False),
        "docx": DocxFmeaExporter().render(snapshot, draft_preview=False),
    }
    assert MARKER not in published["json"].decode("utf-8")
    assert _scan_office_package(published["xlsx"])[1] == {}
    assert _scan_office_package(published["docx"])[1] == {}


def test_office_packages_scan_every_member_and_only_contain_generated_preview_marker() -> None:
    snapshot = make_normalized_snapshot()
    for exporter, expected_hits in (
        (XlsxFmeaExporter(), {"xl/worksheets/sheet3.xml": 1}),
        (DocxFmeaExporter(), {"word/document.xml": 2}),
    ):
        published = exporter.render(snapshot, draft_preview=False)
        preview = exporter.render(snapshot, draft_preview=True)

        published_names, published_hits = _scan_office_package(published)
        preview_names, preview_hits = _scan_office_package(preview)

        assert published_names
        assert preview_names
        assert published_hits == {}
        assert preview_hits == expected_hits


@pytest.mark.parametrize(
    "location",
    (
        "rows",
        "risk_records",
        "propagation",
        "evidence_summary",
        "decision_summary",
        "unresolved_items",
        "version_manifest",
        "audit_summary",
    ),
)
@pytest.mark.parametrize("draft_preview", (False, True), ids=("published", "preview"))
@pytest.mark.parametrize(
    ("exporter_type", "expected_code", "expected_message"),
    (
        (CanonicalJsonExporter, "FMEA_EXPORT_JSON_INVALID", "snapshot cannot be rendered as canonical JSON"),
        (XlsxFmeaExporter, "FMEA_EXPORT_XLSX_INVALID", "snapshot cannot be rendered as XLSX"),
        (DocxFmeaExporter, "FMEA_EXPORT_DOCX_INVALID", "snapshot cannot be rendered as DOCX"),
    ),
    ids=("json", "xlsx", "docx"),
)
def test_all_formats_reject_reserved_marker_in_nested_snapshot_values(
    location: str,
    draft_preview: bool,
    exporter_type,
    expected_code: str,
    expected_message: str,
) -> None:
    snapshot = _snapshot_with_reserved_marker(location)
    exporter = exporter_type()

    failures = []
    for _ in range(2):
        with pytest.raises(ValueError) as captured:
            exporter.render(snapshot, draft_preview=draft_preview)
        failures.append((captured.value.code, str(captured.value)))
        assert captured.value.__cause__ is None
        assert MARKER not in str(captured.value)
        assert "confidential-before" not in str(captured.value)
        assert "confidential-after" not in str(captured.value)

    assert failures == [(expected_code, expected_message)] * 2


def test_repeated_render_preserves_semantic_identity_and_snapshot() -> None:
    snapshot = make_normalized_snapshot(row_payload={"row_id": "row-1", "unicode": "燃烧室 🔥"})
    before = _semantic_json(snapshot)

    first_xlsx = _parse_xlsx(XlsxFmeaExporter().render(snapshot))
    second_xlsx = _parse_xlsx(XlsxFmeaExporter().render(snapshot))
    first_docx = _parse_docx(DocxFmeaExporter().render(snapshot))
    second_docx = _parse_docx(DocxFmeaExporter().render(snapshot))

    assert _without_format_identity(first_xlsx) == _without_format_identity(second_xlsx)
    assert _without_format_identity(first_xlsx) == _without_format_identity(first_docx)
    assert _without_format_identity(first_docx) == _without_format_identity(second_docx)
    assert _without_format_identity(first_docx) == _without_format_identity(before)
    assert _semantic_json(snapshot) == before
