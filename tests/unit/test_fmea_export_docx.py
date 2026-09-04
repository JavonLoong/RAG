from __future__ import annotations

# ruff: noqa: RUF001
import io
from dataclasses import fields
from zipfile import ZipFile

import pytest
from docx import Document
from docx.oxml.ns import qn

from fmea_application.snapshot_contracts import NormalizedFmeaSnapshot, snapshot_content_hash
from fmea_infrastructure.export_docx import DocxFmeaExporter
from tests.fmea_governance_fixtures import make_normalized_snapshot

MARKER = "DRAFT PREVIEW — NOT PUBLISHED"
SECTIONS = ["FMEA", "Risk", "Propagation", "Evidence", "Decisions", "Unresolved"]


def _forge_snapshot(snapshot: NormalizedFmeaSnapshot, **overrides: object) -> NormalizedFmeaSnapshot:
    forged = object.__new__(NormalizedFmeaSnapshot)
    for field in fields(snapshot):
        object.__setattr__(forged, field.name, getattr(snapshot, field.name))
    for name, value in overrides.items():
        object.__setattr__(forged, name, value)
    return forged


def _paragraph_text(document: Document) -> str:
    return "\n".join(paragraph.text for paragraph in document.paragraphs)


def _manifest_table(document: Document):
    return next(
        table
        for table in document.tables
        if tuple(cell.text for cell in table.rows[0].cells) == ("Key", "Value", "Type")
    )


def test_docx_render_has_title_manifest_tables_sections_and_footer_identity() -> None:
    snapshot = make_normalized_snapshot(rows=2)
    exporter = DocxFmeaExporter()

    payload = exporter.render(snapshot)

    assert type(payload) is bytes
    assert exporter.format == "docx"
    assert exporter.media_type == "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
    document = Document(io.BytesIO(payload))
    text = _paragraph_text(document)
    table_text = "\n".join(cell.text for table in document.tables for row in table.rows for cell in row.cells)
    assert "FMEA Export" in text
    assert all(section in text for section in SECTIONS)
    assert snapshot.revision_id in table_text
    assert snapshot.snapshot_id in table_text
    assert snapshot.publication_id in table_text
    assert snapshot.manifest_id in table_text
    assert snapshot.snapshot_hash in table_text
    assert len(document.tables) >= 7
    footer_text = "\n".join(paragraph.text for paragraph in document.sections[0].footer.paragraphs)
    assert snapshot.revision_id in footer_text
    assert snapshot.snapshot_id in footer_text
    assert snapshot.publication_id in footer_text
    assert snapshot.snapshot_hash in footer_text


def test_docx_reading_body_is_prose_with_full_evidence_and_explicit_canonical_markers() -> None:
    from tests.unit.test_fmea_report_view import _layout, _snapshot

    quote = "滤清器堵塞会造成燃料压力下降，维护记录应核对更换周期。" * 8
    snapshot = _snapshot(
        layout=_layout(
            {
                "field_key": "failure_mode",
                "label": "故障模式",
                "value_type": "string",
                "value_path": ("row", "failure_mode"),
            },
            {
                "field_key": "causes",
                "label": "原因",
                "value_type": "string[]",
                "value_path": ("row", "causes"),
            },
            {
                "field_key": "effects",
                "label": "影响",
                "value_type": "string[]",
                "value_path": ("row", "effects"),
            },
            {
                "field_key": "fuel.pressure_drop",
                "label": "压降",
                "value_type": "decimal",
                "value_path": ("extension_values", "fuel.pressure_drop"),
            },
        ),
        row={
            "failure_mode": "燃料滤清器堵塞",
            "causes": ("杂质积聚", "低温结蜡", "维护周期过长"),
            "effects": ("供油压力下降", "燃烧不稳定"),
            "extension_values": (
                {"field_key": "fuel.pressure_drop", "value_type": "decimal", "value": "48.2000"},
            ),
            "blank": "",
        },
        refs=(
            {
                "evidence_id": "evidence-1",
                "document_id": "manual-1",
                "document_version": "7",
                "content_hash": "c" * 64,
                "evidence_hash": "d" * 64,
                "locator": {"page": 12, "span": 3},
                "quote": quote,
                "source_type": "primary_document",
                "source_trust": "trusted",
            },
        ),
    )

    document = Document(io.BytesIO(DocxFmeaExporter().render(snapshot)))
    all_text = _paragraph_text(document) + "\n" + "\n".join(
        cell.text for table in document.tables for row in table.rows for cell in row.cells
    )
    first_headers = tuple(cell.text for cell in document.tables[0].rows[0].cells)
    reading_paragraphs = _paragraph_text(document)
    east_asia_fonts = [
        run._element.rPr.rFonts.get(qn("w:eastAsia"))
        for paragraph in document.paragraphs
        for run in paragraph.runs
        if run._element.rPr is not None and run._element.rPr.rFonts is not None
    ]

    assert first_headers == ("故障模式", "原因", "影响")
    assert not {"评分摘要", "复核状态", "证据编号"} & set(first_headers)
    assert "FMEA 正文" in all_text
    assert "逐行详情" in all_text
    assert "燃料滤清器堵塞" in all_text
    assert quote in all_text
    assert "fuel.pressure_drop" in all_text
    assert "48.2000" in all_text
    assert "mechanisms：（空列表）" in reading_paragraphs
    assert "risk_assessment：（无）" in reading_paragraphs
    assert "blank：（空字符串）" in reading_paragraphs
    assert "证据 evidence-1.quote：" in reading_paragraphs
    assert "证据 evidence-1.document_id：manual-1" in reading_paragraphs
    assert "复核 decision-1.decision：accepted" in reading_paragraphs
    assert not any('"quote":' in paragraph.text for paragraph in document.paragraphs)
    assert "__types__" not in first_headers
    assert "Canonical table: Manifest" in _paragraph_text(document)
    assert "Canonical table: FMEA" in _paragraph_text(document)
    assert "Microsoft YaHei" in east_asia_fonts
    assert document.paragraphs[0].runs[0].font.underline is False
    assert any(
        run.font.size is not None and run.font.size.pt == 11
        for paragraph in document.paragraphs
        for run in paragraph.runs
        if run.text
    )
    assert any(
        paragraph.paragraph_format.space_after is not None and paragraph.paragraph_format.space_after.pt == 6
        for paragraph in document.paragraphs
        if paragraph.text
    )


def test_docx_formula_prefix_text_is_literal_and_package_has_no_executable_or_external_parts() -> None:
    snapshot = make_normalized_snapshot(
        row_payload={
            "row_id": "row-docx",
            "formula": "=SUM(A1:A2)",
            "plus": "+cmd|' /C calc'!A0",
            "minus": "-1+1",
            "at": "@SUM(1,1)",
        }
    )

    payload = DocxFmeaExporter().render(snapshot)

    document = Document(io.BytesIO(payload))
    text = "\n".join(cell.text for table in document.tables for row in table.rows for cell in row.cells)
    assert "=SUM(A1:A2)" in text
    assert "+cmd|' /C calc'!A0" in text
    with ZipFile(io.BytesIO(payload)) as archive:
        original_names = archive.namelist()
        names = {name.casefold() for name in original_names}
        assert not any("vbaproject" in name or "altchunk" in name for name in names)
        assert not any(name.startswith("word/embeddings/") for name in names)
        assert not any(
            b'targetmode="external"' in archive.read(name).lower() for name in original_names if name.endswith(".rels")
        )
        assert b"altChunk" not in archive.read("word/document.xml")


def test_docx_draft_marker_is_explicit_and_published_output_has_none() -> None:
    snapshot = make_normalized_snapshot()
    exporter = DocxFmeaExporter(draft_preview=True)

    draft = exporter.render(snapshot)
    published = exporter.render(snapshot, draft_preview=False)

    draft_text = _paragraph_text(Document(io.BytesIO(draft)))
    published_text = _paragraph_text(Document(io.BytesIO(published)))
    assert MARKER in draft_text
    assert MARKER not in published_text
    draft_manifest_table = _manifest_table(Document(io.BytesIO(draft)))
    draft_manifest = {
        row.cells[0].text: (row.cells[1].text, row.cells[2].text) for row in draft_manifest_table.rows[1:]
    }
    published_manifest_table = _manifest_table(Document(io.BytesIO(published)))
    published_manifest = {
        row.cells[0].text: (row.cells[1].text, row.cells[2].text) for row in published_manifest_table.rows[1:]
    }
    assert draft_manifest["publication_id"] == ("null", "null")
    assert draft_manifest["source_publication_id"] == (snapshot.publication_id, "str")
    assert published_manifest["publication_id"] == (snapshot.publication_id, "str")
    assert published_manifest["source_publication_id"] == ("null", "null")


@pytest.mark.parametrize("bad_value", ["bad\x01value", float("nan"), float("inf"), object()])
def test_docx_rejects_office_unsafe_or_malformed_snapshot_without_leaking_input(bad_value: object) -> None:
    snapshot = make_normalized_snapshot()
    forged = _forge_snapshot(snapshot, rows=({"row_id": "row-1", "value": bad_value},), row_count=1)

    with pytest.raises(ValueError) as captured:
        DocxFmeaExporter().render(forged)

    assert getattr(captured.value, "code", None) == "FMEA_EXPORT_DOCX_INVALID"
    assert str(captured.value) == "snapshot cannot be rendered as DOCX"
    assert "bad" not in str(captured.value)


@pytest.mark.parametrize(
    "overrides",
    (
        {"row_count": 0},
        {"schema_version": "not-normalized-v99"},
    ),
)
def test_docx_rejects_hash_consistent_snapshot_invariant_bypass(overrides: dict[str, object]) -> None:
    snapshot = make_normalized_snapshot()
    forged = _forge_snapshot(snapshot, **overrides)
    object.__setattr__(forged, "snapshot_hash", snapshot_content_hash(forged))

    with pytest.raises(ValueError) as captured:
        DocxFmeaExporter().render(forged)

    assert getattr(captured.value, "code", None) == "FMEA_EXPORT_DOCX_INVALID"
    assert str(captured.value) == "snapshot cannot be rendered as DOCX"
    assert captured.value.__cause__ is None


def test_docx_rejects_wrong_type_and_does_not_mutate_snapshot() -> None:
    snapshot = make_normalized_snapshot()
    before = repr(snapshot)

    with pytest.raises(ValueError) as captured:
        DocxFmeaExporter().render(object())  # type: ignore[arg-type]

    assert getattr(captured.value, "code", None) == "FMEA_EXPORT_SNAPSHOT_INVALID"
    assert str(captured.value) == "snapshot must be a NormalizedFmeaSnapshot"
    DocxFmeaExporter().render(snapshot)
    assert repr(snapshot) == before
