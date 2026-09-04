from __future__ import annotations

import importlib
import io
import json
import shutil
from copy import deepcopy
from hashlib import sha256

import pytest


def _bundle(tmp_path):
    root = tmp_path / "acceptance"
    root.mkdir()
    payload = b'{"cases":[],"schema_version":"graphrag.fmea.full.acceptance.v1"}\n'
    (root / "evidence.json").write_bytes(payload)
    manifest = {
        "schema_version": "graphrag.fmea.full.acceptance.v1",
        "artifact_id": "full-12345678-1234-4234-8234-123456789abc",
        "cases": ["fuel-combustion"],
        "files": {"evidence.json": {"sha256": sha256(payload).hexdigest(), "size_bytes": len(payload)}},
        "summary": {},
    }
    (root / "manifest.json").write_text(json.dumps(manifest), encoding="utf-8")
    return root, manifest


@pytest.mark.parametrize("mutation", ["changed", "missing", "extra", "traversal", "wrong_size"])
def test_bundle_inventory_rejects_tampered_or_unaccounted_files(tmp_path, mutation):
    verifier = importlib.import_module("scripts.verify_fmea_full_acceptance")
    root, manifest = _bundle(tmp_path)
    if mutation == "changed":
        (root / "evidence.json").write_bytes(b"tampered")
    elif mutation == "missing":
        (root / "evidence.json").unlink()
    elif mutation == "extra":
        (root / "unlisted.txt").write_text("unaccounted", encoding="utf-8")
    elif mutation == "traversal":
        manifest["files"]["../outside.json"] = manifest["files"]["evidence.json"]
        (root / "manifest.json").write_text(json.dumps(manifest), encoding="utf-8")
    else:
        manifest["files"]["evidence.json"]["size_bytes"] += 1
        (root / "manifest.json").write_text(json.dumps(manifest), encoding="utf-8")
    with pytest.raises(verifier.VerificationError):
        verifier.load_bundle(root)


def test_bundle_parser_rejects_duplicate_json_keys(tmp_path):
    verifier = importlib.import_module("scripts.verify_fmea_full_acceptance")
    root, _ = _bundle(tmp_path)
    (root / "manifest.json").write_text('{"files":{},"files":{}}', encoding="utf-8")
    with pytest.raises(verifier.VerificationError, match="DUPLICATE"):
        verifier.load_bundle(root)


def test_valid_file_inventory_is_not_itself_full_acceptance(tmp_path):
    verifier = importlib.import_module("scripts.verify_fmea_full_acceptance")
    root, _ = _bundle(tmp_path)
    manifest, files = verifier.load_bundle(root)
    assert manifest["cases"] == ["fuel-combustion"]
    assert set(files) == {"evidence.json"}
    result = verifier.verify_acceptance_directory(root)
    assert result.passed is False
    assert result.error_code


@pytest.mark.parametrize("format_name", ["xlsx", "docx"])
def test_independent_office_parser_preserves_json_semantics(format_name):
    from fmea_infrastructure.export_docx import DocxFmeaExporter
    from fmea_infrastructure.export_json import CanonicalJsonExporter
    from fmea_infrastructure.export_xlsx import XlsxFmeaExporter
    from tests.fmea_governance_fixtures import make_normalized_snapshot

    verifier = importlib.import_module("scripts.verify_fmea_full_acceptance")
    snapshot = make_normalized_snapshot()
    renderer = {"xlsx": XlsxFmeaExporter, "docx": DocxFmeaExporter}[format_name]()
    actual = verifier.parse_export(renderer.render(snapshot), format_name)
    expected = json.loads(CanonicalJsonExporter().render(snapshot))
    assert actual["format"] == format_name
    assert actual["row_count"] == 1
    assert {key: value for key, value in actual.items() if key not in {"format", "media_type"}} == {
        key: value for key, value in expected.items() if key not in {"format", "media_type"}
    }


def _legacy_xlsx_payload(payload: bytes) -> bytes:
    import openpyxl

    workbook = openpyxl.load_workbook(io.BytesIO(payload))
    del workbook["正文详情"]
    del workbook["正文"]
    assert tuple(workbook.sheetnames) == (
        "Manifest",
        "FMEA",
        "Risk",
        "Propagation",
        "Evidence",
        "Decisions",
        "Unresolved",
    )
    output = io.BytesIO()
    workbook.save(output)
    workbook.close()
    return output.getvalue()


def _legacy_docx_payload(payload: bytes) -> bytes:
    from docx import Document
    from docx.text.paragraph import Paragraph

    document = Document(io.BytesIO(payload))
    body = document.element.body
    children = list(body)
    machine_appendix = next(
        index
        for index, child in enumerate(children)
        if child.tag.endswith("}p") and Paragraph(child, document).text.strip() == "机器附录"
    )
    for child in children[:machine_appendix]:
        body.remove(child)
    for child in list(body):
        if child.tag.endswith("}p"):
            body.remove(child)
    assert len(document.tables) == 7
    output = io.BytesIO()
    document.save(output)
    return output.getvalue()


def test_legacy_v1_snapshot_exports_use_the_documented_old_office_layout_portably():
    from core_domain.fmea.governance import canonical_json_value
    from fmea_infrastructure.export_docx import DocxFmeaExporter
    from fmea_infrastructure.export_json import CanonicalJsonExporter
    from fmea_infrastructure.export_xlsx import XlsxFmeaExporter
    from tests.fmea_governance_fixtures import make_normalized_snapshot

    verifier = importlib.import_module("scripts.verify_fmea_full_acceptance")
    snapshot = make_normalized_snapshot()
    assert "body_schema_version" not in snapshot.version_manifest
    payloads = {}
    exports = []
    for renderer in (CanonicalJsonExporter(), XlsxFmeaExporter(), DocxFmeaExporter()):
        path = f"exports/legacy.{renderer.format}"
        payload = renderer.render(snapshot)
        if renderer.format == "xlsx":
            payload = _legacy_xlsx_payload(payload)
        elif renderer.format == "docx":
            payload = _legacy_docx_payload(payload)
        payloads[path] = payload
        exports.append({"path": path, "format": renderer.format})

    case = {"snapshots": [canonical_json_value(snapshot)], "exports": exports}

    assert verifier.verify_export_set(
        case,
        payloads,
        contract_version=verifier.LEGACY_SCHEMA_VERSION,
    ) == set(payloads)


def test_independent_parser_rejects_private_marker_even_in_office_metadata():
    import io
    from zipfile import ZipFile

    from fmea_infrastructure.export_xlsx import XlsxFmeaExporter
    from tests.fmea_governance_fixtures import make_normalized_snapshot

    verifier = importlib.import_module("scripts.verify_fmea_full_acceptance")
    original = XlsxFmeaExporter().render(make_normalized_snapshot())
    output = io.BytesIO()
    with ZipFile(io.BytesIO(original)) as source, ZipFile(output, "w") as target:
        for member in source.infolist():
            content = source.read(member)
            if member.filename == "docProps/core.xml":
                content = content.replace(b"</cp:coreProperties>", b"<secret>Bearer secret-token</secret></cp:coreProperties>")
            target.writestr(member, content)
    with pytest.raises(verifier.VerificationError):
        verifier.parse_export(output.getvalue(), "xlsx")


@pytest.mark.parametrize(
    ("collection", "record", "counter"),
    [
        ("audits", {"command": "fmea.approval.decide", "actor_type": "model"}, "model_approval_count"),
        ("candidates", {"row_id": "row-1", "field_claims": [{"field_key": "failure_mode", "claim_status": "known", "evidence_ids": ["absent"]}]}, "known_without_evidence_count"),
        ("risk_records", {"status": "confirmed", "dimensions": [], "derived": None}, "confirmed_invalid_score_count"),
        ("propagation_graphs", {"edges": [{"risk_priority": "high", "review_status": "accepted", "evidence_ids": []}]}, "accepted_high_risk_evidence_free_edge_count"),
    ],
)
def test_p0_counters_are_derived_from_raw_evidence_not_claimed_summary(collection, record, counter):
    verifier = importlib.import_module("scripts.verify_fmea_full_acceptance")
    case = {collection: [record], "summary": {counter: 0}}
    counts = verifier.count_p0_violations(case)
    assert counts[counter] == 1
    assert sum(counts.values()) == 1


def test_native_snapshot_hash_is_recomputed_after_file_hashes_could_be_rewritten():
    from core_domain.fmea.governance import canonical_json_value
    from tests.fmea_governance_fixtures import make_normalized_snapshot

    verifier = importlib.import_module("scripts.verify_fmea_full_acceptance")
    snapshot = canonical_json_value(make_normalized_snapshot())
    verifier.verify_native_hashes({"snapshots": [snapshot]})
    snapshot["rows"][0]["failure_mode"] = "changed after publication"
    with pytest.raises(verifier.VerificationError, match="SNAPSHOT_HASH"):
        verifier.verify_native_hashes({"snapshots": [snapshot]})


def test_score_counter_uses_native_rule_identity_and_version():
    verifier = importlib.import_module("scripts.verify_fmea_full_acceptance")
    case = {
        "evidence_packs": [{"pack_id": "ep-1", "refs": [{"evidence_id": "ev-1"}]}],
        "scoring_rules": [{"rule_pack_id": "sod", "version": "1", "score_min": 1, "score_max": 10}],
        "risk_records": [{
            "status": "confirmed", "rule_pack_id": "sod", "rule_pack_version": "1", "evidence_pack_id": "ep-1",
            "dimensions": [{"name": name, "value": value, "evidence_ids": ["ev-1"]} for name, value in (("severity", 5), ("occurrence", 3), ("detection", 2))],
            "derived": {"decision_severity": 5, "occurrence": 3, "detection": 2, "rpn": 30},
        }],
    }
    assert verifier.count_p0_violations(case)["confirmed_invalid_score_count"] == 0
    case["risk_records"][0]["rule_pack_version"] = "stale"
    assert verifier.count_p0_violations(case)["confirmed_invalid_score_count"] == 1


@pytest.mark.parametrize("mutation", [None, "row", "hash", "missing", "duplicate", "unbound"])
def test_export_set_binds_every_format_to_the_actual_snapshot(mutation):
    from core_domain.fmea.governance import canonical_json_value
    from fmea_infrastructure.export_docx import DocxFmeaExporter
    from fmea_infrastructure.export_json import CanonicalJsonExporter
    from fmea_infrastructure.export_xlsx import XlsxFmeaExporter
    from tests.fmea_governance_fixtures import make_normalized_snapshot

    verifier = importlib.import_module("scripts.verify_fmea_full_acceptance")
    snapshot = make_normalized_snapshot()
    payloads = {}
    exports = []
    for renderer in (CanonicalJsonExporter(), XlsxFmeaExporter(), DocxFmeaExporter()):
        path = f"exports/fuel.{renderer.format}"
        payloads[path] = renderer.render(snapshot)
        exports.append({"path": path, "format": renderer.format})
    case = {"snapshots": [canonical_json_value(snapshot)], "exports": exports}
    if mutation in {"row", "hash", "unbound"}:
        edited = json.loads(payloads["exports/fuel.json"])
        if mutation == "row":
            edited["rows"][0]["failure_mode"] = "tampered export"
        elif mutation == "hash":
            edited["snapshot_hash"] = "0" * 64
        else:
            edited["snapshot_id"] = "unbound-snapshot"
        payloads["exports/fuel.json"] = json.dumps(edited).encode()
    elif mutation == "missing":
        exports.pop()
    elif mutation == "duplicate":
        exports.append(exports[0])
    if mutation is None:
        assert verifier.verify_export_set(case, payloads) == set(payloads)
    else:
        with pytest.raises(verifier.VerificationError):
            verifier.verify_export_set(case, payloads)


@pytest.fixture(scope="module")
def executed_governance_records(tmp_path_factory):
    # This is evidence for the governance validator only, not an assertion that
    # the older governance runner exercised candidate generation or scoring.
    from scripts.run_fmea_governance_acceptance import run_acceptance

    result = run_acceptance(output_root=tmp_path_factory.mktemp("governance"))
    case = {}
    for name in ("revisions", "approvals", "publications", "snapshots", "audits", "outbox"):
        case[name] = json.loads((result.artifact_dir / f"{name}.json").read_bytes())["items"]
    case["audits"] = [item["event"] for item in case["audits"]]
    return case


@pytest.mark.parametrize("mutation", [None, "approval_hash", "actor", "audit_actor", "snapshot_publication", "duplicate_event"])
def test_publication_chain_binds_revision_approval_snapshot_and_human_audit(executed_governance_records, mutation):
    from copy import deepcopy

    verifier = importlib.import_module("scripts.verify_fmea_full_acceptance")
    case = deepcopy(executed_governance_records)
    if mutation == "approval_hash":
        case["approvals"][0]["revision_hash"] = "0" * 64
    elif mutation == "actor":
        case["publications"][0]["publisher_actor_id"] = "unbound-actor"
    elif mutation == "audit_actor":
        next(item for item in case["audits"] if item["command"] == "fmea.approval.decide")["actor_type"] = "model"
    elif mutation == "snapshot_publication":
        case["snapshots"][0]["publication_id"] = "unbound-publication"
    elif mutation == "duplicate_event":
        case["audits"].append(case["audits"][0])
    if mutation is None:
        verifier.verify_publication_bindings(case)
    else:
        with pytest.raises(verifier.VerificationError):
            verifier.verify_publication_bindings(case)


@pytest.mark.parametrize("format_name", ["xlsx", "docx"])
def test_office_identity_column_cannot_disagree_with_row_id(format_name):
    import io

    import openpyxl
    from docx import Document

    from fmea_infrastructure.export_docx import DocxFmeaExporter
    from fmea_infrastructure.export_xlsx import XlsxFmeaExporter
    from tests.fmea_governance_fixtures import make_normalized_snapshot

    verifier = importlib.import_module("scripts.verify_fmea_full_acceptance")
    snapshot = make_normalized_snapshot()
    output = io.BytesIO()
    if format_name == "xlsx":
        document = openpyxl.load_workbook(io.BytesIO(XlsxFmeaExporter().render(snapshot)))
        document["FMEA"]["A2"] = "forged-display-identity"
    else:
        document = Document(io.BytesIO(DocxFmeaExporter().render(snapshot)))
        fmea_table = next(
            table
            for table in document.tables
            if table.rows and table.rows[0].cells and table.rows[0].cells[0].text == "Identity"
            and any(cell.text == "row_id" for cell in table.rows[0].cells)
        )
        fmea_table.cell(1, 0).text = "forged-display-identity"
    document.save(output)
    with pytest.raises(verifier.VerificationError, match="IDENTITY"):
        verifier.parse_export(output.getvalue(), format_name)


def test_incomplete_full_verifier_cli_fails_closed(tmp_path, capsys, monkeypatch):
    verifier = importlib.import_module("scripts.verify_fmea_full_acceptance")
    root, _ = _bundle(tmp_path)
    assert verifier.main([str(root)]) == 1
    result = json.loads(capsys.readouterr().out)
    assert result["passed"] is False
    assert result["error_code"] == "FMEA_WORKFLOW_EVIDENCE_INCOMPLETE"
    # The developer may already have a valid public latest artifact. Test an
    # isolated missing pointer through the real resolver, never their output.
    resolve_latest = verifier.resolve_latest_directory
    monkeypatch.setattr(verifier, "resolve_latest_directory", lambda: resolve_latest(tmp_path / "missing-latest"))
    assert verifier.main(["--latest"]) == 1
    assert json.loads(capsys.readouterr().out)["passed"] is False


def test_extension_claim_cannot_hide_unsupported_canonical_fields():
    verifier = importlib.import_module("scripts.verify_fmea_full_acceptance")
    case = {
        "evidence_packs": [{"pack_id": "ep", "refs": [{"evidence_id": "ev"}]}],
        "candidates": [{
            "row_id": "r", "record_version": 1, "evidence_pack_id": "ep", "claim_status": "known",
            "failure_mode": "unsupported failure", "field_evidence": [],
            "field_claims": [{"field_key": "electrical.voltage", "claim_status": "known", "evidence_ids": ["ev"]}],
        }],
    }
    assert verifier.count_p0_violations(case)["known_without_evidence_count"] == 1


def test_full_evidence_rejects_duplicate_cases_before_workflow_claims(tmp_path):
    verifier = importlib.import_module("scripts.verify_fmea_full_acceptance")
    root = tmp_path / "acceptance"
    root.mkdir()
    evidence = {
        "schema_version": verifier.SCHEMA_VERSION,
        "cases": [
            {"case_id": "fuel-combustion", "coverage": "full_lifecycle"},
            {"case_id": "fuel-combustion", "coverage": "full_lifecycle"},
        ],
    }
    payload = json.dumps(evidence, separators=(",", ":")).encode()
    (root / "evidence.json").write_bytes(payload)
    manifest = {
        "schema_version": verifier.SCHEMA_VERSION,
        "artifact_id": "3f6f4f6f-4e1a-4b2f-9e6f-3f6f4f6f4e1a",
        "cases": ["fuel-combustion", "fuel-combustion"],
        "summary": {},
        "files": {"evidence.json": {"sha256": sha256(payload).hexdigest(), "size_bytes": len(payload)}},
    }
    (root / "manifest.json").write_text(json.dumps(manifest), encoding="utf-8")

    result = verifier.verify_acceptance_directory(root)

    assert result.passed is False
    assert result.error_code == "FMEA_DUPLICATE_CASE"


def test_latest_resolves_only_safe_uuid_pointer_and_artifact_directory(tmp_path):
    verifier = importlib.import_module("scripts.verify_fmea_full_acceptance")
    reports = tmp_path / "fmea-full-acceptance"
    reports.mkdir()
    artifact_id = "3f6f4f6f-4e1a-4b2f-9e6f-3f6f4f6f4e1a"
    target = reports / artifact_id
    target.mkdir()
    (reports / "latest.json").write_text(json.dumps({"artifact_id": artifact_id}), encoding="utf-8")

    assert verifier.resolve_latest_directory(reports) == target

    (reports / "latest.json").write_text(json.dumps({"artifact_id": "../outside"}), encoding="utf-8")
    with pytest.raises(verifier.VerificationError, match="LATEST_POINTER_INVALID"):
        verifier.resolve_latest_directory(reports)


def test_workflow_semantics_reject_invented_steps_without_raw_categories():
    verifier = importlib.import_module("scripts.verify_fmea_full_acceptance")
    case = {
        "case_id": "fuel-combustion",
        "coverage": "full_lifecycle",
        "steps": [{"command": "evidence.select", "result_ids": {"pack_id": "invented"}}],
        "summary": {
            "model_approval_count": 0,
            "known_without_evidence_count": 0,
            "confirmed_invalid_score_count": 0,
            "accepted_high_risk_evidence_free_edge_count": 0,
        },
    }

    with pytest.raises(verifier.VerificationError, match="WORKFLOW_EVIDENCE_INCOMPLETE"):
        verifier.validate_case_semantics(case, case["summary"])


def test_acceptance_requires_one_full_fuel_case_not_structural_only(tmp_path):
    verifier = importlib.import_module("scripts.verify_fmea_full_acceptance")
    root, manifest = _bundle(tmp_path)
    evidence = {
        "schema_version": manifest["schema_version"],
        "cases": [{"case_id": "electrical-demo", "coverage": "structural_domain"}],
    }
    payload = json.dumps(evidence, separators=(",", ":")).encode()
    (root / "evidence.json").write_bytes(payload)
    manifest["cases"] = ["electrical-demo"]
    manifest["summary"] = dict.fromkeys(verifier._P0_FIELDS, 0)
    manifest["files"]["evidence.json"] = {"sha256": sha256(payload).hexdigest(), "size_bytes": len(payload)}
    (root / "manifest.json").write_text(json.dumps(manifest), encoding="utf-8")

    result = verifier.verify_acceptance_directory(root)

    assert result.passed is False
    assert result.error_code == "FMEA_WORKFLOW_EVIDENCE_INCOMPLETE"


def test_acceptance_rejects_payload_outside_exact_used_file_set():
    verifier = importlib.import_module("scripts.verify_fmea_full_acceptance")
    cases = [{
        "case_id": "fuel-combustion",
        "coverage": "full_lifecycle",
        "exports": [{"path": "exports/fuel.json"}],
        "template_import_sources": [{"path": "inputs/template.xlsx"}],
    }]

    with pytest.raises(verifier.VerificationError, match="ARTIFACT_UNUSED_FILE"):
        verifier._validate_used_payloads(
            cases,
            {
                "evidence.json": b"",
                "exports/fuel.json": b"",
                "inputs/template.xlsx": b"",
                "inputs/unconsumed.xlsx": b"",
            },
        )


@pytest.fixture(scope="module")
def full_artifact(tmp_path_factory):
    """Execute one real offline lifecycle per module on every machine/CI."""
    from scripts.run_fmea_full_acceptance import run_full_acceptance

    verifier = importlib.import_module("scripts.verify_fmea_full_acceptance")
    artifact = run_full_acceptance(output_root=tmp_path_factory.mktemp("fmea-security-acceptance")).artifact_dir
    manifest, payloads = verifier.load_bundle(artifact)
    evidence = verifier._parse(payloads["evidence.json"])
    return verifier, artifact, manifest, payloads, evidence


def test_real_full_artifact_passes_independent_verifier(full_artifact, tmp_path_factory):
    verifier, artifact, manifest, _payloads, evidence = full_artifact

    result = verifier.verify_acceptance_directory(artifact)

    assert result.passed is True
    assert artifact.is_relative_to(tmp_path_factory.getbasetemp())
    assert result.artifact_id == manifest["artifact_id"]
    case = evidence["cases"][0]
    assert case["coverage"] == "full_lifecycle"
    propagation_replays = [item for item in case["replays"] if item["command"].startswith("fmea.propagation.")]
    assert propagation_replays
    assert all(item.get("state_hash_before") == item.get("state_hash_after") for item in propagation_replays)


def _rewrite_office_payload(payload, format_name, mutation, old_hash, new_hash):  # noqa: C901 - explicit format tamper matrix
    if format_name == "xlsx":
        import openpyxl

        workbook = openpyxl.load_workbook(io.BytesIO(payload))
        if mutation == "failure_body":
            for sheet in workbook.worksheets:
                for row in sheet.iter_rows():
                    for cell in row:
                        if cell.value == "fuel filter blockage":
                            cell.value = "forged failure body"
        elif mutation == "quote":
            for sheet in workbook.worksheets:
                for row in sheet.iter_rows():
                    for cell in row:
                        if isinstance(cell.value, str):
                            cell.value = cell.value.replace("Synthetic acceptance fixture", "Tampered acceptance fixture")
        elif mutation == "scoring_association":
            for sheet in workbook.worksheets:
                for row in sheet.iter_rows():
                    for cell in row:
                        if cell.value == "fuel-sod-rpn":
                            cell.value = "forged-sod-rpn"
        elif mutation == "review_version":
            table = workbook["Decisions"]
            headers = [cell.value for cell in next(table.iter_rows())]
            column = headers.index("record_version") + 1
            table.cell(row=2, column=column).value = 3
        elif mutation == "missing_body_marker":
            table = workbook["Manifest"]
            for row in table.iter_rows(min_row=2):
                if row[0].value == "version_manifest":
                    version_manifest = json.loads(row[1].value)
                    version_manifest.pop("body_schema_version", None)
                    row[1].value = json.dumps(version_manifest, ensure_ascii=False, sort_keys=True, separators=(",", ":"))
                    break
        elif mutation == "layout_label":
            table = workbook["Manifest"]
            for row in table.iter_rows(min_row=2):
                if row[0].value == "version_manifest":
                    version_manifest = json.loads(row[1].value)
                    version_manifest["report_layout"]["columns"][0]["label"] = "forged visible label"
                    row[1].value = json.dumps(version_manifest, ensure_ascii=False, sort_keys=True, separators=(",", ":"))
                    break
        elif mutation == "report_identity_switch":
            table = workbook["Manifest"]
            for row in table.iter_rows(min_row=2):
                if row[0].value == "version_manifest":
                    version_manifest = json.loads(row[1].value)
                    _switch_report_identity(version_manifest)
                    row[1].value = json.dumps(version_manifest, ensure_ascii=False, sort_keys=True, separators=(",", ":"))
                    break
        output = io.BytesIO()
        workbook.save(output)
        workbook.close()
        return output.getvalue().replace(old_hash.encode(), new_hash.encode())

    from docx import Document
    from docx.table import Table
    from docx.text.paragraph import Paragraph

    document = Document(io.BytesIO(payload))
    canonical = {}
    marker = None
    for child in document.element.body.iterchildren():
        if child.tag.endswith("}p"):
            text = Paragraph(child, document).text.strip()
            if text.startswith("Canonical table: "):
                marker = text.removeprefix("Canonical table: ")
        elif child.tag.endswith("}tbl") and marker is not None:
            canonical[marker] = Table(child, document)
            marker = None
    if mutation == "review_version":
        table = canonical["Decisions"]
        headers = [cell.text for cell in table.rows[0].cells]
        table.cell(1, headers.index("record_version")).text = "3"
    elif mutation == "missing_body_marker":
        table = canonical["Manifest"]
        for row in table.rows[1:]:
            if row.cells[0].text == "version_manifest":
                version_manifest = json.loads(row.cells[1].text)
                version_manifest.pop("body_schema_version", None)
                row.cells[1].text = json.dumps(version_manifest, ensure_ascii=False, sort_keys=True, separators=(",", ":"))
                break
    elif mutation == "layout_label":
        table = canonical["Manifest"]
        for row in table.rows[1:]:
            if row.cells[0].text == "version_manifest":
                version_manifest = json.loads(row.cells[1].text)
                version_manifest["report_layout"]["columns"][0]["label"] = "forged visible label"
                row.cells[1].text = json.dumps(version_manifest, ensure_ascii=False, sort_keys=True, separators=(",", ":"))
                break
    elif mutation == "report_identity_switch":
        table = canonical["Manifest"]
        for row in table.rows[1:]:
            if row.cells[0].text == "version_manifest":
                version_manifest = json.loads(row.cells[1].text)
                _switch_report_identity(version_manifest)
                row.cells[1].text = json.dumps(version_manifest, ensure_ascii=False, sort_keys=True, separators=(",", ":"))
                break
    else:
        old_value, new_value = {
            "failure_body": ("fuel filter blockage", "forged failure body"),
            "quote": ("Synthetic acceptance fixture", "Tampered acceptance fixture"),
            "scoring_association": ("fuel-sod-rpn", "forged-sod-rpn"),
        }[mutation]
        for table in document.tables:
            for row in table.rows:
                for cell in row.cells:
                    if old_value in cell.text:
                        for paragraph in cell.paragraphs:
                            for run in paragraph.runs:
                                run.text = run.text.replace(old_value, new_value)
                        if old_value in cell.text:
                            cell.text = cell.text.replace(old_value, new_value)
    output = io.BytesIO()
    document.save(output)
    return output.getvalue().replace(old_hash.encode(), new_hash.encode())


def _switch_report_identity(version_manifest):
    identities = version_manifest["template_identities"]
    current = version_manifest["report_layout"]["template_identity"]
    current_identity = [current["template_id"], current["version"], current["template_hash"]]
    alternate = next((item for item in identities if item != current_identity), None)
    assert alternate is not None, "fixture must contain a second approved template identity"
    version_manifest["report_layout"]["template_identity"] = {
        "template_id": alternate[0],
        "version": alternate[1],
        "template_hash": alternate[2],
    }


def _reseal_body_tamper(full_artifact, tmp_path, mutation):  # noqa: C901 - explicit body tamper matrix
    verifier, artifact, manifest, _payloads, evidence = full_artifact
    root = tmp_path / manifest["artifact_id"]
    shutil.copytree(artifact, root)
    evidence = deepcopy(evidence)
    original_case = evidence["cases"][0]
    native = {name: deepcopy(original_case[name]) for name in ("candidates", "risk_records", "evidence_packs", "review_decisions", "propagation_graphs", "scoring_rules")}
    old_hashes = {}
    new_hashes = {}
    for snapshot in original_case["snapshots"]:
        old_hash = snapshot["snapshot_hash"]
        old_hashes[snapshot["snapshot_id"]] = old_hash
        if mutation == "failure_body":
            snapshot["rows"][0]["failure_mode"] = "forged failure body"
        elif mutation == "quote":
            snapshot["evidence_summary"][0]["refs"][0]["quote"] = snapshot["evidence_summary"][0]["refs"][0]["quote"].replace(
                "Synthetic acceptance fixture", "Tampered acceptance fixture"
            )
        elif mutation == "scoring_association":
            risk = snapshot["risk_records"][0]
            risk["rule_pack_id"] = "forged-sod-rpn"
            risk["derived"]["scoring_rule_pack_id"] = "forged-sod-rpn"
        elif mutation == "review_version":
            snapshot["decision_summary"][0]["record_version"] = 3
        elif mutation == "missing_body_marker":
            snapshot["version_manifest"].pop("body_schema_version")
        elif mutation == "layout_label":
            snapshot["version_manifest"]["report_layout"]["columns"][0]["label"] = "forged visible label"
        elif mutation == "report_identity_switch":
            _switch_report_identity(snapshot["version_manifest"])
        else:  # pragma: no cover - parametrized below
            raise AssertionError(mutation)
        new_hash = _adversarial_digest({key: value for key, value in snapshot.items() if key != "snapshot_hash"})
        snapshot["snapshot_hash"] = new_hash
        new_hashes[snapshot["snapshot_id"]] = new_hash

    case = evidence["cases"][0]
    changed_payloads = {}
    for export in case["exports"]:
        snapshot_id = export["run"]["snapshot_id"]
        path = export["path"]
        format_name = export["format"]
        payload = (root / path).read_bytes()
        old_hash = old_hashes[snapshot_id]
        new_hash = new_hashes[snapshot_id]
        if format_name == "json":
            view = json.loads(payload)
            if mutation == "failure_body":
                view["rows"][0]["failure_mode"] = "forged failure body"
            elif mutation == "quote":
                view["evidence_summary"][0]["refs"][0]["quote"] = view["evidence_summary"][0]["refs"][0]["quote"].replace(
                    "Synthetic acceptance fixture", "Tampered acceptance fixture"
                )
            elif mutation == "scoring_association":
                risk = view["risk_records"][0]
                risk["rule_pack_id"] = "forged-sod-rpn"
                risk["derived"]["scoring_rule_pack_id"] = "forged-sod-rpn"
            elif mutation == "review_version":
                view["decision_summary"][0]["record_version"] = 3
            elif mutation == "missing_body_marker":
                view["version_manifest"].pop("body_schema_version")
            elif mutation == "layout_label":
                view["version_manifest"]["report_layout"]["columns"][0]["label"] = "forged visible label"
            elif mutation == "report_identity_switch":
                _switch_report_identity(view["version_manifest"])
            view["snapshot_hash"] = new_hash
            payload = json.dumps(view, ensure_ascii=False, sort_keys=True, separators=(",", ":")).encode()
        else:
            payload = _rewrite_office_payload(payload, format_name, mutation, old_hash, new_hash)
        (root / path).write_bytes(payload)
        changed_payloads[path] = payload
        export["run"]["snapshot_hash"] = new_hash
        export["manifest"]["snapshot_hash"] = new_hash
        export["manifest"]["sha256"] = sha256(payload).hexdigest()
        export["manifest"]["byte_length"] = len(payload)

    evidence_payload = json.dumps(evidence, ensure_ascii=False, sort_keys=True, separators=(",", ":")).encode()
    (root / "evidence.json").write_bytes(evidence_payload)
    manifest = deepcopy(manifest)
    for path, payload in {**changed_payloads, "evidence.json": evidence_payload}.items():
        manifest["files"][path] = {"sha256": sha256(payload).hexdigest(), "size_bytes": len(payload)}
    (root / "manifest.json").write_text(json.dumps(manifest, sort_keys=True), encoding="utf-8")
    assert {name: original_case[name] for name in native} == native
    return verifier, root, native


def _rewrite_visible_body(payload, format_name, surface):  # noqa: C901 - explicit XLSX/DOCX bypass matrix
    replacements = 0
    if format_name == "xlsx":
        import openpyxl

        workbook = openpyxl.load_workbook(io.BytesIO(payload))
        sheet_name = "正文" if surface == "main" else "正文详情"
        sheet = workbook[sheet_name]
        old_value = "fuel filter blockage" if surface == "main" else "Synthetic acceptance fixture"
        new_value = "forged visible body" if surface == "main" else "forged visible evidence"
        for row in sheet.iter_rows():
            for cell in row:
                if isinstance(cell.value, str) and old_value in cell.value:
                    cell.value = cell.value.replace(old_value, new_value)
                    replacements += 1
        output = io.BytesIO()
        workbook.save(output)
        workbook.close()
        return output.getvalue(), replacements

    from docx import Document

    document = Document(io.BytesIO(payload))
    old_value = "fuel filter blockage" if surface == "main" else "Synthetic acceptance fixture"
    new_value = "forged visible body" if surface == "main" else "forged visible evidence"
    if surface == "main":
        cells = (cell for row in document.tables[0].rows for cell in row.cells)
    else:
        for paragraph in document.paragraphs:
            if old_value in paragraph.text:
                paragraph.text = paragraph.text.replace(old_value, new_value)
                replacements += 1
    if surface == "main":
        for cell in cells:
            if old_value in cell.text:
                cell.text = cell.text.replace(old_value, new_value)
                replacements += 1
    output = io.BytesIO()
    document.save(output)
    return output.getvalue(), replacements


def _reseal_visible_body_tamper(full_artifact, tmp_path, format_name, surface):
    verifier, artifact, manifest, _payloads, evidence = full_artifact
    root = tmp_path / manifest["artifact_id"]
    shutil.copytree(artifact, root)
    evidence = deepcopy(evidence)
    changed_payloads = {}
    replacement_count = 0
    expected_publications = sum(1 for export in evidence["cases"][0]["exports"] if export["format"] == format_name)
    assert expected_publications == 2
    for export in evidence["cases"][0]["exports"]:
        path = export["path"]
        current_format = export["format"]
        payload = (root / path).read_bytes()
        if current_format == format_name:
            assert current_format != "json"
            payload, replacements = _rewrite_visible_body(payload, current_format, surface)
            replacement_count += replacements
            (root / path).write_bytes(payload)
            export["manifest"]["sha256"] = sha256(payload).hexdigest()
            export["manifest"]["byte_length"] = len(payload)
        changed_payloads[path] = payload

    assert replacement_count == expected_publications and replacement_count > 0

    evidence_payload = json.dumps(evidence, ensure_ascii=False, sort_keys=True, separators=(",", ":")).encode()
    (root / "evidence.json").write_bytes(evidence_payload)
    manifest = deepcopy(manifest)
    for path, payload in {**changed_payloads, "evidence.json": evidence_payload}.items():
        manifest["files"][path] = {"sha256": sha256(payload).hexdigest(), "size_bytes": len(payload)}
    (root / "manifest.json").write_text(json.dumps(manifest, sort_keys=True), encoding="utf-8")
    return verifier, root


@pytest.mark.parametrize("surface", ["main", "details"])
@pytest.mark.parametrize("format_name", ["xlsx", "docx"])
def test_v2_visible_body_tampering_is_rejected_after_resealing_office_exports(full_artifact, tmp_path, format_name, surface):
    verifier, root = _reseal_visible_body_tamper(full_artifact, tmp_path, format_name, surface)

    result = verifier.verify_acceptance_directory(root)

    assert result.passed is False
    assert result.error_code == "FMEA_PUBLICATION_VISIBLE_BODY_MISMATCH"


@pytest.mark.parametrize(
    ("mutation", "error_code"),
    [
        ("failure_body", "FMEA_PUBLICATION_BODY_NATIVE_MISMATCH"),
        ("quote", "FMEA_PUBLICATION_BODY_NATIVE_MISMATCH"),
        ("scoring_association", "FMEA_PUBLICATION_BODY_NATIVE_MISMATCH"),
        ("review_version", "FMEA_PUBLICATION_BODY_NATIVE_MISMATCH"),
        ("missing_body_marker", "FMEA_PUBLICATION_BODY_MARKER_MISSING"),
    ],
)
def test_v2_body_tampering_is_rejected_after_resealing_all_exports(full_artifact, tmp_path, mutation, error_code):
    verifier, root, _native = _reseal_body_tamper(full_artifact, tmp_path, mutation)

    result = verifier.verify_acceptance_directory(root)

    assert result.passed is False
    assert result.error_code == error_code


@pytest.mark.parametrize("mutation", ["layout_label", "report_identity_switch"])
def test_v2_layout_tampering_is_rejected_after_resealing_all_exports(full_artifact, tmp_path, mutation):
    verifier, root, _native = _reseal_body_tamper(full_artifact, tmp_path, mutation)

    result = verifier.verify_acceptance_directory(root)

    assert result.passed is False
    assert result.error_code == "FMEA_PUBLICATION_LAYOUT_MISMATCH"


@pytest.mark.parametrize(
    ("mutation", "error_code"),
    [
        ("review_row", "FMEA_REVIEW_BINDING_INVALID"),
        ("source_row_hash", "FMEA_PROPAGATION_BINDING_INVALID"),
        ("missing_lineage", "FMEA_PROPAGATION_LINEAGE_INVALID"),
        ("lineage_hash", "FMEA_PROPAGATION_LINEAGE_INVALID"),
        ("duplicate_risk_history", "FMEA_DUPLICATE_RECORD"),
        ("duplicate_graph_id", "FMEA_DUPLICATE_RECORD"),
        ("replay_flag", "FMEA_REPLAY_SCHEMA_INVALID"),
        ("replay_state_hash", "FMEA_REPLAY_STATE_HASH_MISMATCH"),
        ("lifecycle_binding", "FMEA_LIFECYCLE_BINDING_INVALID"),
        ("export_bytes", "FMEA_EXPORT_HASH_MISMATCH"),
    ],
)
def test_real_full_artifact_tamper_matrix(full_artifact, mutation, error_code):  # noqa: C901 - explicit independent tamper cases
    verifier, _artifact, manifest, original_payloads, evidence = full_artifact
    case = deepcopy(evidence["cases"][0])
    payloads = dict(original_payloads)

    if mutation == "review_row":
        case["review_decisions"][0]["row"]["failure_mode"] = "forged review row"
    elif mutation == "source_row_hash":
        case["source_row_bindings"][0]["row_hash"] = "0" * 64
    elif mutation == "missing_lineage":
        case.pop("source_row_lineage")
    elif mutation == "lineage_hash":
        case["source_row_lineage"][0]["canonical_row_hash"] = "0" * 64
    elif mutation == "duplicate_risk_history":
        case["risk_records"].append(deepcopy(case["risk_records"][0]))
    elif mutation == "duplicate_graph_id":
        case["propagation_graphs"].append(deepcopy(case["propagation_graphs"][0]))
    elif mutation == "replay_flag":
        replay = next(item for item in case["replays"] if "replayed" in item["first"])
        replay["first"]["replayed"] = True
    elif mutation == "replay_state_hash":
        replay = next(item for item in case["replays"] if item["command"].startswith("fmea.propagation."))
        replay["state_hash_after"] = "0" * 64
    elif mutation == "lifecycle_binding":
        event = next(item for item in case["lifecycle_events"] if item.get("supersession_id"))
        event["old_publication_id"] = "forged-publication"
    elif mutation == "export_bytes":
        path = next(iter(case["exports"]))["path"]
        payloads[path] += b"\n"

    with pytest.raises(verifier.VerificationError, match=error_code):
        verifier.validate_case_semantics(case, manifest["summary"], payloads)


def _adversarial_digest(value):
    """Rehash tampered public DTOs independently of verifier/domain helpers."""
    return sha256(json.dumps(value, ensure_ascii=False, sort_keys=True, separators=(",", ":"), allow_nan=False).encode()).hexdigest()


def _replace_receipt_value(value, original, replacement):
    if value == original:
        return deepcopy(replacement)
    if isinstance(value, dict):
        return {key: _replace_receipt_value(item, original, replacement) for key, item in value.items()}
    if isinstance(value, list):
        return [_replace_receipt_value(item, original, replacement) for item in value]
    return value


def _forge_local_dto(case, kind):
    collection = "candidates" if kind == "row" else "risk_records"
    original = deepcopy(next(item for item in case[collection] if kind == "row" or item["status"] == "confirmed"))
    forged = deepcopy(original)
    if kind == "row":
        forged["failure_mode"] = "fuel filter rupture"
    else:
        severity = next(item for item in forged["dimensions"] if item["name"] == "severity")
        severity["value"] = 7 if severity["value"] != 7 else 6
        derived = forged["derived"]
        derived["decision_severity"] = severity["value"]
        derived["severity_by_consequence_class"] = [["decision", severity["value"]]]
        derived["rpn"] = severity["value"] * derived["occurrence"] * derived["detection"]
    for name in (collection, "review_decisions", "steps", "replays"):
        case[name] = _replace_receipt_value(case[name], original, forged)
    if kind == "row":
        digest = "sha256:" + _adversarial_digest(forged)
        for binding in case["source_row_bindings"]:
            if binding["row_id"] == forged["row_id"]:
                binding["row_hash"] = binding["persisted_row_hash_after"] = digest
        for lineage in case["source_row_lineage"]:
            if lineage["source_row_id"] == forged["row_id"]:
                lineage["canonical_row_hash"] = digest


def _forge_migration_target(case, field):
    original = deepcopy(case["migration_reports"][0])
    forged = deepcopy(original)
    if field == "pack_hash":
        forged["target_domain_pack_identity"][2] = "0" * 64
    else:
        forged["target_revision_hash"] = "0" * 64
    forged["report_hash"] = _adversarial_digest({key: value for key, value in forged.items() if key not in {"report_hash", "created_at"}})
    assert forged["report_hash"] != original["report_hash"]
    case["migration_reports"][0] = forged
    for name in ("migration_results", "steps", "replays"):
        case[name] = _replace_receipt_value(case[name], original, forged)
        case[name] = _replace_receipt_value(case[name], original["report_hash"], forged["report_hash"])


@pytest.mark.parametrize(
    "mutation",
    [
        "row_local_dto",
        "risk_local_dto",
        "approval_empty_replay",
        "migration_pack_hash",
        "migration_target_revision_hash",
        "graph_unknown_target",
        "export_run_snapshot_hash",
        "export_run_publication_id",
        "export_run_workspace_id",
        "export_run_revision_id",
        "lifecycle_self_supersession",
        "lifecycle_status",
        "lifecycle_outbox",
    ],
)
def test_real_full_artifact_review_binding_gaps(full_artifact, mutation):  # noqa: C901 - six independent review findings
    verifier, _artifact, manifest, payloads, evidence = full_artifact
    original = evidence["cases"][0]
    case = deepcopy(original)

    if mutation in {"row_local_dto", "risk_local_dto"}:
        _forge_local_dto(case, mutation.split("_", 1)[0])
    elif mutation == "approval_empty_replay":
        replay = next(item for item in case["replays"] if item["command"] == "fmea.approval.decide")
        replay["first"] = {}
        replay["replayed"] = {}
    elif mutation.startswith("migration_"):
        _forge_migration_target(case, "pack_hash" if mutation == "migration_pack_hash" else "target_revision_hash")
    elif mutation == "graph_unknown_target":
        graph = next(item for item in case["propagation_graphs"] if item["status"] == "confirmed")
        missing = "absent-topology-node"
        assert all(node["node_id"] != missing for topology in case["topology_snapshots"] for node in topology["nodes"])
        graph["edges"][0]["target_entity_id"] = missing
    elif mutation.startswith("export_run_"):
        field = mutation.removeprefix("export_run_")
        case["exports"][0]["run"][field] = "0" * 64 if field == "snapshot_hash" else "forged-resource"
        assert case["exports"][0]["manifest"] == original["exports"][0]["manifest"]
    elif mutation == "lifecycle_self_supersession":
        event = next(item for item in case["lifecycle_events"] if "supersession_id" in item)
        event["old_publication_id"] = event["new_publication_id"]
    elif mutation == "lifecycle_status":
        lifecycle = next(item for item in case["publication_lifecycle"] if item["effective_status"] == "superseded")
        lifecycle["effective_status"] = "published"
    elif mutation == "lifecycle_outbox":
        event = next(item for item in case["outbox"] if item["event_type"] == "publication.withdrawn")
        event["payload"]["withdrawal"]["reason"] = "forged withdrawal reason"
        event["payload_hash"] = "sha256:" + _adversarial_digest(event["payload"])
        audit = next(item for item in case["audits"] if item["command"] == "fmea.publication.withdraw" and item["row_id"] == event["aggregate_id"])
        audit["canonical_payload_hash"] = event["payload_hash"]

    # These immutable witnesses deliberately retain the actual published data.
    for name in ("revisions", "snapshots", "publications"):
        assert case[name] == original[name]
    if mutation != "lifecycle_outbox":
        assert case["outbox"] == original["outbox"]
    assert case != original
    with pytest.raises(verifier.VerificationError):
        verifier.validate_case_semantics(case, manifest["summary"], payloads)
