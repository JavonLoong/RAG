"""Presentation-only DOCX rendering for normalized FMEA snapshots."""

# ruff: noqa: RUF001

from __future__ import annotations

import io
from collections.abc import Mapping, Sequence
from typing import Final, NoReturn
from xml.etree.ElementTree import ParseError
from zipfile import BadZipFile, ZipFile

import orjson
from defusedxml.ElementTree import fromstring as safe_xml_fromstring
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor

from fmea_application.report_view import build_report_view
from fmea_application.snapshot_contracts import (
    DRAFT_PREVIEW_MARKER,
    NormalizedFmeaSnapshot,
    revalidate_normalized_snapshot,
)

from .export_json import _snapshot_projection, _validate_export_value

_MAX_DOCX_CELL_TEXT: Final = 1_000_000
_MAX_COLUMNS: Final = 256
_TYPES_COLUMN: Final = "__types__"
_RESERVED_HEADERS: Final = frozenset({"Identity", _TYPES_COLUMN})
_REL_NS: Final = "http://schemas.openxmlformats.org/package/2006/relationships"
_READABLE_MAIN_MAX_COLUMNS: Final = 3
_EAST_ASIA_FONT: Final = "Microsoft YaHei"


class DocxExportError(ValueError):
    """Stable, public-safe DOCX rendering failure."""

    def __init__(self, code: str, message: str) -> None:
        self.code = code
        super().__init__(message)


def _error(code: str, message: str) -> DocxExportError:
    return DocxExportError(code, message)


def _invalid() -> NoReturn:
    raise ValueError


def _is_xml_char(value: str) -> bool:
    return all(
        codepoint in {0x9, 0xA, 0xD}
        or 0x20 <= codepoint <= 0xD7FF
        or 0xE000 <= codepoint <= 0xFFFD
        or 0x10000 <= codepoint <= 0x10FFFF
        for codepoint in map(ord, value)
    )


def _validate_office_value(value: object) -> None:
    if isinstance(value, str):
        if not _is_xml_char(value):
            _invalid()
        return
    if isinstance(value, Mapping):
        for key, item in value.items():
            if not isinstance(key, str) or not _is_xml_char(key):
                _invalid()
            _validate_office_value(item)
        return
    if isinstance(value, Sequence) and not isinstance(value, str | bytes):
        for item in value:
            _validate_office_value(item)


def _json_text(value: object) -> str:
    try:
        return orjson.dumps(value, option=orjson.OPT_SORT_KEYS).decode("utf-8")
    except (TypeError, ValueError, OverflowError):
        _invalid()


def _value_encoding(value: object) -> tuple[str, str]:
    if type(value) is str:
        return value, "str"
    if value is None:
        return "null", "null"
    if type(value) is bool:
        return "true" if value else "false", "bool"
    if type(value) is int:
        return str(value), "int"
    if type(value) is float:
        return _json_text(value), "float"
    if isinstance(value, Mapping | list):
        return _json_text(value), "json"
    _invalid()


def _cell_text(value: str) -> str:
    if len(value) > _MAX_DOCX_CELL_TEXT:
        _invalid()
    return value


def _set_run_font(run) -> None:
    run.font.name = _EAST_ASIA_FONT
    run._element.get_or_add_rPr().get_or_add_rFonts().set(qn("w:eastAsia"), _EAST_ASIA_FONT)


def _set_cell_text(cell, value: str, *, font_size: int = 9) -> None:
    cell.text = _cell_text(value)
    for paragraph in cell.paragraphs:
        paragraph.paragraph_format.space_after = Pt(0)
        for run in paragraph.runs:
            run.font.size = Pt(font_size)
            _set_run_font(run)


def _identity(record: Mapping[str, object], identity_field: str | None, index: int) -> str:
    if identity_field is not None:
        value = record.get(identity_field)
        if type(value) is not str or not value.strip():
            _invalid()
        return value
    for field_name in ("source_id", "code", "item_id"):
        value = record.get(field_name)
        if type(value) is str and value.strip():
            return value
    return f"item-{index:03d}"


def _record_columns(records: Sequence[Mapping[str, object]]) -> list[str]:
    keys = {key for record in records for key in record}
    if any(type(key) is not str or key in _RESERVED_HEADERS for key in keys):
        _invalid()
    columns = sorted(keys)
    if len(columns) + 2 > _MAX_COLUMNS:
        _invalid()
    return columns


def _add_heading(document: Document, text: str, *, level: int):
    paragraph = document.add_heading(text, level=level)
    paragraph_properties = paragraph._p.get_or_add_pPr()
    paragraph_border = paragraph_properties.find(qn("w:pBdr"))
    if paragraph_border is not None:
        paragraph_properties.remove(paragraph_border)
    for run in paragraph.runs:
        run.font.color.rgb = RGBColor(0, 0, 0)
        run.font.underline = False
        _set_run_font(run)
    return paragraph


def _remove_style_paragraph_borders(document: Document) -> None:
    for style in document.styles:
        style_properties = style._element.find(qn("w:pPr"))
        if style_properties is None:
            continue
        paragraph_border = style_properties.find(qn("w:pBdr"))
        if paragraph_border is not None:
            style_properties.remove(paragraph_border)


def _append_manifest(document: Document, projection: Mapping[str, object]) -> None:
    _add_heading(document, "Canonical table: Manifest", level=1)
    table = document.add_table(rows=1, cols=3)
    table.style = "Table Grid"
    table.autofit = False
    for column, width in zip(table.columns, (1.2, 4.3, 0.7), strict=True):
        column.width = Inches(width)
        for cell in column.cells:
            cell.width = Inches(width)
    for cell, value in zip(table.rows[0].cells, ("Key", "Value", "Type"), strict=True):
        _set_cell_text(cell, value)
    metadata: tuple[tuple[str, object], ...] = (
        ("schema_version", projection["schema_version"]),
        ("snapshot_schema_version", projection["snapshot_schema_version"]),
        ("snapshot_id", projection["snapshot_id"]),
        ("workspace_id", projection["workspace_id"]),
        ("analysis_id", projection["analysis_id"]),
        ("revision_id", projection["revision_id"]),
        ("revision_hash", projection["revision_hash"]),
        ("publication_id", projection["publication_id"]),
        ("source_publication_id", projection["source_publication_id"]),
        ("manifest_id", projection["manifest_id"]),
        ("row_count", projection["row_count"]),
        ("risk_count", len(projection["risk_records"])),
        ("propagation_present", projection["propagation"] is not None),
        ("evidence_count", len(projection["evidence_summary"])),
        ("decision_count", len(projection["decision_summary"])),
        ("unresolved_count", len(projection["unresolved_items"])),
        ("snapshot_hash", projection["snapshot_hash"]),
        ("created_at", projection["created_at"]),
        ("version_manifest", projection["version_manifest"]),
        ("audit_summary", projection["audit_summary"]),
        ("draft_preview", projection["draft_preview"]),
        ("draft_marker", projection["draft_marker"]),
        ("format", projection["format"]),
        ("media_type", projection["media_type"]),
    )
    for key, value in metadata:
        encoded, value_type = _value_encoding(value)
        row = table.add_row().cells
        _set_cell_text(row[0], key)
        _set_cell_text(row[1], encoded)
        _set_cell_text(row[2], value_type)


def _append_typed_table(
    document: Document,
    records: Sequence[Mapping[str, object]],
    *,
    identity_field: str | None,
) -> None:
    columns = _record_columns(records)
    table = document.add_table(rows=1, cols=len(columns) + 2)
    table.style = "Table Grid"
    headers = ["Identity", *columns, _TYPES_COLUMN]
    for cell, header in zip(table.rows[0].cells, headers, strict=True):
        _set_cell_text(cell, header)
    for row_index, record in enumerate(records, start=2):
        row = table.add_row().cells
        _set_cell_text(row[0], _identity(record, identity_field, row_index - 1))
        types: dict[str, str] = {}
        for column_index, key in enumerate(columns, start=1):
            if key not in record:
                _set_cell_text(row[column_index], "")
                continue
            encoded, value_type = _value_encoding(record[key])
            types[key] = value_type
            _set_cell_text(row[column_index], encoded)
        _set_cell_text(row[-1], _json_text(types))


def _plain_value(value: object) -> object:
    if isinstance(value, Mapping):
        return {str(key): _plain_value(item) for key, item in value.items()}
    if isinstance(value, Sequence) and not isinstance(value, str | bytes):
        return [_plain_value(item) for item in value]
    return value


def _human_text(value: object) -> str:
    if value is None:
        return "（无）"
    if isinstance(value, str):
        if value == "":
            return "（空字符串）"
        if not value.strip():
            return "（空白字符串）"
        return value
    if isinstance(value, Mapping):
        if not value:
            return "（空对象）"
        return _json_text(_plain_value(value))
    if isinstance(value, Sequence) and not isinstance(value, str | bytes):
        if not value:
            return "（空列表）"
        return "；".join(_human_text(item) for item in value)
    return str(value)


def _readable_columns(view) -> tuple[object, ...]:
    return tuple(view.columns[:_READABLE_MAIN_MAX_COLUMNS])


def _detail_row(detail: Mapping[str, object]) -> Mapping[str, object]:
    row = detail.get("row")
    return row if isinstance(row, Mapping) else detail


def _detail_risks(detail: Mapping[str, object], row: Mapping[str, object]) -> tuple[Mapping[str, object], ...]:
    row_id = row.get("row_id")
    record_version = row.get("record_version")
    records = detail.get("risk_records", ())
    if not isinstance(records, Sequence) or isinstance(records, str | bytes):
        return ()
    return tuple(
        risk
        for risk in records
        if isinstance(risk, Mapping)
        and risk.get("row_id") == row_id
        and risk.get("source_record_version") == record_version
    )


def _detail_evidence_ids(row: Mapping[str, object]) -> tuple[str, ...]:
    result: set[str] = set()
    bindings = row.get("field_evidence", ())
    if isinstance(bindings, Sequence) and not isinstance(bindings, str | bytes):
        for binding in bindings:
            if not isinstance(binding, Mapping):
                continue
            evidence_ids = binding.get("evidence_ids", ())
            if isinstance(evidence_ids, Sequence) and not isinstance(evidence_ids, str | bytes):
                result.update(str(evidence_id) for evidence_id in evidence_ids)
    return tuple(sorted(result))


def _append_readable_main(document: Document, view, *, title: str) -> None:
    _add_heading(document, title, level=1)
    columns = _readable_columns(view)
    headers = [column.label for column in columns]
    table = document.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    for cell, header in zip(table.rows[0].cells, headers, strict=True):
        _set_cell_text(cell, str(header), font_size=11)
    for values in view.rows:
        display_values = [_human_text(values[column.field_key]) for column in columns]
        cells = table.add_row().cells
        for cell, value in zip(cells, display_values, strict=True):
            _set_cell_text(cell, value, font_size=11)


def _append_readable_paragraph(document: Document, text: str) -> None:
    paragraph = document.add_paragraph()
    paragraph.paragraph_format.space_after = Pt(6)
    run = paragraph.add_run(_cell_text(text))
    run.font.size = Pt(11)
    _set_run_font(run)


def _append_mapping_paragraphs(document: Document, prefix: str, value: object) -> None:
    if isinstance(value, Mapping):
        if not value:
            _append_readable_paragraph(document, f"{prefix}：（空对象）")
            return
        for key, item in value.items():
            field = f"{prefix}.{key}" if prefix else str(key)
            _append_mapping_paragraphs(document, field, item)
        return
    if (
        isinstance(value, Sequence)
        and not isinstance(value, str | bytes)
        and value
        and all(isinstance(item, Mapping) for item in value)
    ):
        for index, item in enumerate(value, start=1):
            _append_mapping_paragraphs(document, f"{prefix}[{index}]", item)
        return
    _append_readable_paragraph(document, f"{prefix}：{_human_text(value)}")


def _append_row_details(document: Document, view) -> None:  # noqa: C901
    _add_heading(document, "逐行详情", level=1)
    for row_index, detail in enumerate(view.details):
        row = _detail_row(detail)
        row_id = _human_text(row.get("row_id"))
        record_version = _human_text(row.get("record_version"))
        _add_heading(document, f"行ID：{row_id}（记录版本：{record_version}）", level=2)
        declared_keys = set()
        values = view.rows[row_index]
        for column in view.columns:
            declared_keys.add(column.field_key)
            _append_readable_paragraph(
                document,
                f"{column.label} [{column.field_key}]：{_human_text(values[column.field_key])}",
            )
        for key, value in row.items():
            if key == "extension_values" or key in declared_keys:
                continue
            _append_readable_paragraph(document, f"{key}：{_human_text(value)}")
        extension_values = row.get("extension_values", ())
        if isinstance(extension_values, Sequence) and not isinstance(extension_values, str | bytes):
            for extension in extension_values:
                if isinstance(extension, Mapping):
                    _append_readable_paragraph(
                        document,
                        f"扩展字段 {extension.get('field_key')}（{extension.get('value_type')}）："
                        f"{_human_text(extension.get('value'))}",
                    )
        for risk in _detail_risks(detail, row):
            _append_mapping_paragraphs(document, f"评分 {risk.get('assessment_id')}", risk)
        for decision in detail.get("decision_summary", ()):
            if isinstance(decision, Mapping):
                _append_mapping_paragraphs(document, f"复核 {decision.get('decision_id')}", decision)


def _append_global_details(document: Document, view, projection: Mapping[str, object]) -> None:  # noqa: C901
    seen_packs: set[str] = set()
    seen_evidence: set[object] = set()
    summaries = [detail.get("evidence_summary", ()) for detail in view.details]
    if not summaries:
        summaries = [projection["evidence_summary"]]
    _add_heading(document, "共享证据", level=1)
    for summary in summaries:
        if not isinstance(summary, Sequence) or isinstance(summary, str | bytes):
            continue
        for pack in summary:
            if not isinstance(pack, Mapping):
                continue
            pack_id = _human_text(pack.get("pack_id"))
            if pack_id not in seen_packs:
                seen_packs.add(pack_id)
                pack_without_refs = {key: value for key, value in pack.items() if key != "refs"}
                _append_mapping_paragraphs(document, f"证据包 {pack_id}", pack_without_refs)
            refs = pack.get("refs", ())
            if not isinstance(refs, Sequence) or isinstance(refs, str | bytes):
                continue
            for reference in refs:
                if not isinstance(reference, Mapping):
                    continue
                identity = reference.get("evidence_id")
                if not isinstance(identity, str):
                    identity = _json_text(_plain_value(reference))
                if identity in seen_evidence:
                    continue
                seen_evidence.add(identity)
                _append_mapping_paragraphs(document, f"证据 {identity}", reference)
    if not view.details:
        for decision in projection["decision_summary"]:
            if isinstance(decision, Mapping):
                _append_mapping_paragraphs(document, f"复核 {decision.get('decision_id')}", decision)
    if projection["propagation"] is not None:
        _add_heading(document, "传播", level=1)
        _append_mapping_paragraphs(document, "传播", projection["propagation"])


def _append_section(
    document: Document,
    title: str,
    records: Sequence[Mapping[str, object]],
    *,
    identity_field: str | None,
) -> None:
    _add_heading(document, f"Canonical table: {title}", level=1)
    _append_typed_table(document, records, identity_field=identity_field)


def _validate_package_xml(name: str, raw: bytes) -> None:
    folded_name = name.casefold()
    if not folded_name.endswith((".xml", ".rels")):
        return
    root = safe_xml_fromstring(raw, forbid_dtd=True, forbid_entities=True, forbid_external=True)
    if folded_name.endswith(".rels"):
        for relationship in root.iter("{http://schemas.openxmlformats.org/package/2006/relationships}Relationship"):
            target = relationship.attrib.get("Target", "")
            if relationship.attrib.get("TargetMode", "").casefold() == "external" or target.casefold().startswith((
                "http:",
                "https:",
                "file:",
                "ftp:",
            )):
                _invalid()
    if folded_name.endswith(".xml"):
        for element in root.iter():
            if element.tag.rsplit("}", 1)[-1].casefold() in {"altchunk", "fldsimple", "fldchar", "instrtext"}:
                _invalid()


def _validate_package(payload: bytes) -> None:
    try:
        with ZipFile(io.BytesIO(payload)) as archive:
            names = tuple(archive.namelist())
            folded = {name.casefold() for name in names}
            if any(".." in name.split("/") or "\\" in name or name.startswith("/") for name in names):
                _invalid()
            if any("vbaproject" in name or name.startswith("word/embeddings/") for name in folded):
                _invalid()
            for name in names:
                _validate_package_xml(name, archive.read(name))
    except (BadZipFile, OSError, ValueError, ParseError):
        _invalid()


class DocxFmeaExporter:
    """Render a normalized snapshot to an in-memory, presentation-only DOCX."""

    format = "docx"
    media_type = "application/vnd.openxmlformats-officedocument.wordprocessingml.document"

    def __init__(self, draft_preview: bool = False) -> None:
        if type(draft_preview) is not bool:
            raise _error("FMEA_EXPORT_DOCX_INVALID", "draft_preview must be a boolean")
        self._draft_preview = draft_preview

    def render(self, snapshot: NormalizedFmeaSnapshot, *, draft_preview: bool | None = None) -> bytes:
        if type(snapshot) is not NormalizedFmeaSnapshot:
            raise _error("FMEA_EXPORT_SNAPSHOT_INVALID", "snapshot must be a NormalizedFmeaSnapshot")
        if draft_preview is not None and type(draft_preview) is not bool:
            raise _error("FMEA_EXPORT_DOCX_INVALID", "draft_preview must be a boolean or None")
        try:
            resolved_preview = self._draft_preview if draft_preview is None else draft_preview
            snapshot = revalidate_normalized_snapshot(snapshot)
            projection = _snapshot_projection(
                snapshot,
                draft_preview=resolved_preview,
                export_format=self.format,
                media_type=self.media_type,
            )
            _validate_export_value(projection)
            _validate_office_value(projection)
            report_view = build_report_view(snapshot)
            document = Document()
            _remove_style_paragraph_borders(document)
            document.core_properties.title = "FMEA Export"
            title = _add_heading(document, "FMEA Export", level=0)
            title.alignment = WD_ALIGN_PARAGRAPH.CENTER
            if resolved_preview:
                marker = document.add_paragraph(DRAFT_PREVIEW_MARKER)
                marker.alignment = WD_ALIGN_PARAGRAPH.CENTER
            readable_title = (
                "FMEA 正文"
                if projection["version_manifest"].get("body_schema_version") is not None
                else "FMEA 摘要（旧快照兼容视图）"
            )
            _append_readable_main(document, report_view, title=readable_title)
            if report_view.details:
                _append_row_details(document, report_view)
            else:
                _add_heading(document, "逐行详情", level=1)
                _append_readable_paragraph(document, "旧快照仅包含摘要视图，未声明可用的逐行正文。")
            _append_global_details(document, report_view, projection)
            _add_heading(document, "机器附录", level=1)
            _append_manifest(document, projection)
            _append_section(document, "FMEA", projection["rows"], identity_field="row_id")
            _append_section(document, "Risk", projection["risk_records"], identity_field="assessment_id")
            propagation = () if projection["propagation"] is None else (projection["propagation"],)
            _append_section(document, "Propagation", propagation, identity_field=None)
            _append_section(document, "Evidence", projection["evidence_summary"], identity_field="pack_id")
            _append_section(document, "Decisions", projection["decision_summary"], identity_field="decision_id")
            _append_section(document, "Unresolved", projection["unresolved_items"], identity_field=None)
            footer = document.sections[0].footer.paragraphs[0]
            publication_id = projection["publication_id"] or ""
            source_publication_id = projection["source_publication_id"] or ""
            footer.text = (
                f"revision_id={projection['revision_id']} | snapshot_id={projection['snapshot_id']} | "
                f"publication_id={publication_id} | source_publication_id={source_publication_id} | "
                f"snapshot_hash={projection['snapshot_hash']}"
            )
            footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
            buffer = io.BytesIO()
            document.save(buffer)
            payload = buffer.getvalue()
            _validate_package(payload)
        except DocxExportError:
            raise
        except Exception:
            raise _error("FMEA_EXPORT_DOCX_INVALID", "snapshot cannot be rendered as DOCX") from None
        else:
            return payload


__all__ = ["DocxExportError", "DocxFmeaExporter"]
