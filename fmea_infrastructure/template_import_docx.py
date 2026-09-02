"""Bounded, non-executing DOCX template ingestion."""

# ZIP and XML safety checks happen before python-docx opens the package.
# TRY003 is consistent with the stable ValueError-style importer boundary.

from __future__ import annotations

import io
from collections.abc import Callable
from datetime import datetime, timezone
from hashlib import sha256
from uuid import NAMESPACE_URL, uuid5

from docx import Document

from core_domain.fmea.filename_policy import validate_filename
from core_domain.fmea.template_migration import SourceStructureItem, TemplateDraft, TemplateDraftStatus

from .template_import_excel import (
    OfficePackageLimits,
    TemplateImportError,
    _error,
    _parse_xml,
    _text,
    classify_source_fields,
    inspect_office_zip,
)

_REL_NAMESPACE = "http://schemas.openxmlformats.org/package/2006/relationships"
_EXECUTABLE_MARKERS = ("altchunk", "oleobject", "object", "embeddedpackage", "embeddedobject")


def _clock_now() -> str:
    return datetime.now(timezone.utc).isoformat(timespec="microseconds").replace("+00:00", "Z")


def _stable_draft_id(workspace_id: str, source_hash: str) -> str:
    return f"draft-{uuid5(NAMESPACE_URL, f'fmea-template:{workspace_id}:{source_hash}')}"


def _relationship_count(raw: bytes) -> int:
    root = _parse_xml(raw, label="relationship part")
    return len(root.findall(f"{{{_REL_NAMESPACE}}}Relationship"))


def _local_name(tag: object) -> str:
    return tag.rsplit("}", 1)[-1] if isinstance(tag, str) else ""


class DocxTemplateImporter:
    """Read paragraphs and tables while refusing fields, links, and executable parts."""

    def __init__(
        self,
        *,
        limits: OfficePackageLimits | None = None,
        clock: Callable[[], str] = _clock_now,
    ) -> None:
        self._limits = limits or OfficePackageLimits()
        self._clock = clock

    def parse(self, raw_bytes: bytes, filename: str, *, workspace_id: str) -> TemplateDraft:  # noqa: C901
        try:
            validate_filename(filename, "source_filename", expected_extension="docx")
        except ValueError as exc:
            raise _error("FMEA_TEMPLATE_IMPORT_INVALID", "source filename is invalid") from exc
        workspace = _text(workspace_id, "workspace_id")
        package = inspect_office_zip(raw_bytes, filename, kind="docx", limits=self._limits)
        document_xml = package.members["word/document.xml"]
        root = _parse_xml(document_xml, label="Word document")
        if any(_local_name(element.tag).casefold() in _EXECUTABLE_MARKERS for element in root.iter()):
            raise _error("FMEA_TEMPLATE_EXECUTABLE_CONTENT", "Word executable content is unsupported")
        if any(tag in document_xml.lower() for tag in (b"fldsimple", b"fldchar", b"instrtext", b"doctext")):
            raise _error("FMEA_TEMPLATE_FIELD_UNSUPPORTED", "Word fields are unsupported")
        relationship_total = sum(
            _relationship_count(payload)
            for name, payload in package.members.items()
            if name.casefold().endswith(".rels")
        )
        if relationship_total > self._limits.max_relationships:
            raise _error("FMEA_TEMPLATE_LIMIT_EXCEEDED", "Word document has too many relationships")

        try:
            document = Document(io.BytesIO(raw_bytes))
        except Exception as exc:
            raise _error("FMEA_TEMPLATE_CONTAINER_INVALID", "Word document cannot be parsed") from exc

        structure: list[SourceStructureItem] = []
        headers: list[tuple[str, str]] = []
        if len(document.paragraphs) > self._limits.max_paragraphs:
            raise _error("FMEA_TEMPLATE_LIMIT_EXCEEDED", "Word document has too many paragraphs")
        if len(document.tables) > self._limits.max_tables:
            raise _error("FMEA_TEMPLATE_LIMIT_EXCEEDED", "Word document has too many tables")
        for index, paragraph in enumerate(document.paragraphs):
            text = paragraph.text.strip()
            if not text:
                continue
            if len(text) > self._limits.max_text_length:
                raise _error("FMEA_TEMPLATE_LIMIT_EXCEEDED", "paragraph text exceeds the configured limit")
            locator = f"document#paragraph-{index}"
            structure.append(SourceStructureItem(kind="paragraph", locator=locator, value=text))
            headers.append((text, locator))
        for table_index, table in enumerate(document.tables):
            for row_index, row in enumerate(table.rows):
                for cell_index, cell in enumerate(row.cells):
                    text = cell.text.strip()
                    if not text:
                        continue
                    if len(text) > self._limits.max_text_length:
                        raise _error("FMEA_TEMPLATE_LIMIT_EXCEEDED", "table text exceeds the configured limit")
                    locator = f"document#table-{table_index}/row-{row_index}/cell-{cell_index}"
                    structure.append(SourceStructureItem(kind="table-cell", locator=locator, value=text))
                    headers.append((text, locator))
        for name, payload in package.members.items():
            if not name.casefold().endswith(".rels"):
                continue
            rel_root = _parse_xml(payload, label="relationship part")
            for relationship in rel_root.findall(f"{{{_REL_NAMESPACE}}}Relationship"):
                relationship_id = relationship.attrib.get("Id", "")
                if relationship_id:
                    structure.append(SourceStructureItem(kind="relationship", locator=f"{name}#{relationship_id}"))
        if len(structure) > self._limits.max_structure_items:
            raise _error("FMEA_TEMPLATE_LIMIT_EXCEEDED", "Word document structure exceeds the configured limit")
        proposed, unknown, ambiguous, identified = classify_source_fields(tuple(headers))
        source_hash = sha256(raw_bytes).hexdigest()
        return TemplateDraft(
            draft_id=_stable_draft_id(workspace, source_hash),
            workspace_id=workspace,
            source_filename=validate_filename(filename, "source_filename", expected_extension="docx"),
            source_sha256=source_hash,
            source_type="docx",
            structure=tuple(structure),
            proposed_fields=proposed,
            unknown_fields=unknown,
            ambiguous_fields=ambiguous,
            parser_warnings=(),
            status=TemplateDraftStatus.DRAFT,
            created_at=self._clock(),
            identified_fields=identified,
        )


__all__ = ["DocxTemplateImporter", "OfficePackageLimits", "TemplateImportError"]
